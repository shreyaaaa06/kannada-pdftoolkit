import os
import zipfile
from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
import fitz
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import config
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
import html
import platform
import subprocess
import shutil
from pathlib import Path
import io

class PDFOperations:
    def __init__(self):
        self.config = config.Config()
    
    def merge_pdfs(self, file_paths, session_id):
        """Merge multiple PDF files"""
        try:
            if not file_paths or len(file_paths) < 2:
                raise Exception("ವಿಲೀನಕ್ಕೆ ಕನಿಷ್ಠ 2 PDF ಫೈಲ್‌ಗಳು ಅಗತ್ಯ")
        
            valid_files = []
            for file_path in file_paths:
                if not os.path.exists(file_path):
                    continue
                
                if os.path.getsize(file_path) == 0:
                    continue
            
                try:
                    reader = PdfReader(file_path)
                    if len(reader.pages) > 0:
                        valid_files.append(file_path)
                except Exception:
                    continue
        
            if len(valid_files) < 2:
                raise Exception("ವಿಲೀನಕ್ಕೆ ಕನಿಷ್ಠ 2 ಸರಿಯಾದ PDF ಫೈಲ್‌ಗಳು ಅಗತ್ಯ")
            
            writer = PdfWriter()
            
            for file_path in valid_files:
                try:
                    reader = PdfReader(file_path)
                    for page in reader.pages:
                        writer.add_page(page)
                except Exception as e:
                    continue
            
            output_filename = f"{session_id}_merged.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
                raise Exception("ವಿಲೀನ ಫೈಲ್ ರಚನೆ ವಿಫಲವಾಗಿದೆ")
            
            return output_path
            
        except Exception as e:
            raise Exception(f"PDF ವಿಲೀನ ವಿಫಲ: {str(e)}")

    def split_pdf(self, file_path, session_id, pages=""):
        """Split PDF into separate files"""
        try:
            if not os.path.exists(file_path):
                raise Exception("PDF ಫೈಲ್ ಸಿಗಲಿಲ್ಲ")
            
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            
            if total_pages < 2:
                raise Exception("ವಿಭಜನೆಗೆ ಕನಿಷ್ಠ 2 ಪುಟಗಳು ಬೇಕಾಗುತ್ತವೆ")
            
            if pages:
                page_ranges = self._parse_page_ranges(pages, total_pages)
                
                output_filename = f"{session_id}_split.pdf"
                output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
                
                writer = PdfWriter()
                for page_num in page_ranges:
                    if 1 <= page_num <= total_pages:
                        writer.add_page(reader.pages[page_num - 1])
                
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                
                return output_path
            else:
                split_folder = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_split")
                os.makedirs(split_folder, exist_ok=True)
                
                for i, page in enumerate(reader.pages):
                    writer = PdfWriter()
                    writer.add_page(page)
                    
                    page_filename = f"page_{i+1}.pdf"
                    page_path = os.path.join(split_folder, page_filename)
                    
                    with open(page_path, 'wb') as output_file:
                        writer.write(output_file)
                
                zip_filename = f"{session_id}_split.zip"
                zip_path = os.path.join(self.config.OUTPUT_FOLDER, zip_filename)
                
                with zipfile.ZipFile(zip_path, 'w') as zip_file:
                    for root, dirs, files in os.walk(split_folder):
                        for file in files:
                            file_path = os.path.join(root, file)
                            zip_file.write(file_path, file)
                
                shutil.rmtree(split_folder)
                return zip_path
                
        except Exception as e:
            raise Exception(f"PDF ವಿಭಜನೆ ವಿಫಲ: {str(e)}")

    def _parse_page_ranges(self, pages_str, total_pages):
        """Parse page ranges like '1,3,5-10' into list of page numbers"""
        pages = []
        parts = pages_str.replace(' ', '').split(',')
        
        for part in parts:
            if '-' in part:
                start, end = part.split('-', 1)
                try:
                    start_num = int(start)
                    end_num = int(end)
                    pages.extend(range(start_num, end_num + 1))
                except ValueError:
                    continue
            else:
                try:
                    pages.append(int(part))
                except ValueError:
                    continue
        
        return [p for p in pages if 1 <= p <= total_pages]

    def extract_pages(self, file_path, pages, session_id):
        """Extract specific pages from PDF"""
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            page_numbers = self._parse_page_ranges(pages, total_pages)
            
            if not page_numbers:
                raise Exception("ಸರಿಯಾದ ಪುಟ ಸಂಖ್ಯೆಗಳನ್ನು ನಮೂದಿಸಿ")
            
            writer = PdfWriter()
            for page_num in page_numbers:
                writer.add_page(reader.pages[page_num - 1])
            
            output_filename = f"{session_id}_extracted.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"ಪುಟ ಹೊರತೆಗೆಯುವಿಕೆ ವಿಫಲ: {str(e)}")

    def delete_pages(self, file_path, pages, session_id):
        """Delete specific pages from PDF"""
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            pages_to_delete = set(self._parse_page_ranges(pages, total_pages))
            
            writer = PdfWriter()
            for i, page in enumerate(reader.pages):
                if (i + 1) not in pages_to_delete:
                    writer.add_page(page)
            
            output_filename = f"{session_id}_deleted.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"ಪುಟ ಅಳಿಸುವಿಕೆ ವಿಫಲ: {str(e)}")

    def compress_pdf(self, file_path, compression_level, session_id):
        """Compress PDF file"""
        try:
            output_filename = f"{session_id}_compressed.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            try:
                return self._compress_pymupdf(file_path, output_path, compression_level)
            except Exception:
                return self._compress_pypdf2(file_path, output_path, compression_level)
                
        except Exception as e:
            raise Exception(f"PDF ಸಂಕುಚನ ವಿಫಲ: {str(e)}")

    def _compress_pymupdf(self, input_path, output_path, level):
        """Compress PDF using PyMuPDF"""
        doc = fitz.open(input_path)
        
        deflate_level = {
            'low': 1,
            'medium': 6,
            'high': 9
        }.get(level, 6)
        
        doc.save(output_path, deflate=True, deflate_level=deflate_level, clean=True)
        doc.close()
        
        return output_path

    def _compress_pypdf2(self, input_path, output_path, level):
        """Compress PDF using PyPDF2"""
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            page.compress_content_streams()
            writer.add_page(page)
        
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        return output_path

    def pdf_to_images(self, file_path, session_id):
        """Convert PDF pages to JPEG images"""
        try:
            doc = fitz.open(file_path)
            images_folder = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_images")
            os.makedirs(images_folder, exist_ok=True)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                mat = fitz.Matrix(2, 2)
                pix = page.get_pixmap(matrix=mat)
                
                image_filename = f"page_{page_num + 1}.jpg"
                image_path = os.path.join(images_folder, image_filename)
                pix.save(image_path)
            
            doc.close()
            
            zip_filename = f"{session_id}_images.zip"
            zip_path = os.path.join(self.config.OUTPUT_FOLDER, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w') as zip_file:
                for root, dirs, files in os.walk(images_folder):
                    for file in files:
                        file_path = os.path.join(root, file)
                        zip_file.write(file_path, file)
            
            shutil.rmtree(images_folder)
            return zip_path
            
        except Exception as e:
            raise Exception(f"PDF ಚಿತ್ರ ಪರಿವರ್ತನೆ ವಿಫಲ: {str(e)}")

    def images_to_pdf(self, image_paths, session_id):
        """Convert images to PDF"""
        try:
            output_filename = f"{session_id}_from_images.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            images = []
            for image_path in image_paths:
                if os.path.exists(image_path):
                    img = Image.open(image_path)
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    images.append(img)
            
            if images:
                images[0].save(output_path, save_all=True, append_images=images[1:])
            
            return output_path
            
        except Exception as e:
            raise Exception(f"ಚಿತ್ರ PDF ಪರಿವರ್ತನೆ ವಿಫಲ: {str(e)}")

    def pdf_to_word(self, file_path, session_id):
        """Convert PDF to Word document"""
        try:
            doc = fitz.open(file_path)
            word_doc = Document()
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    word_doc.add_paragraph(text)
                
                if page_num < len(doc) - 1:
                    word_doc.add_page_break()
            
            doc.close()
            
            output_filename = f"{session_id}_converted.docx"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            word_doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"PDF Word ಪರಿವರ್ತನೆ ವಿಫಲ: {str(e)}")

    def word_to_pdf(self, file_path, session_id):
        """Convert Word document to PDF"""
        try:
            return self._simple_word_to_pdf(file_path, session_id)
        except Exception as e:
            raise Exception(f"Word PDF ಪರಿವರ್ತನೆ ವಿಫಲ: {str(e)}")

    def _simple_word_to_pdf(self, file_path, session_id):
        """Simple Word to PDF conversion"""
        doc = Document(file_path)
        
        output_filename = f"{session_id}_from_word.pdf"
        output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
        
        pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                p = Paragraph(html.escape(para.text), styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 12))
        
        pdf_doc.build(story)
        return output_path

    def sort_pdf_by_page_numbers(self, file_path, session_id, pages=""):
        """Sort PDF pages by detected Kannada page numbers"""
        try:
            from .kannada_numeral_converter import KannadaNumeralConverter
            
            converter = KannadaNumeralConverter()
            doc = fitz.open(file_path)
            
            page_data = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                extracted_number = converter.extract_page_number_from_text(text)
                
                page_data.append({
                    'page': page,
                    'original_num': page_num + 1,
                    'extracted_num': extracted_number if extracted_number else page_num + 1
                })
            
            page_data.sort(key=lambda x: x['extracted_num'])
            
            output_filename = f"{session_id}_sortedbynum.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            new_doc = fitz.open()
            for data in page_data:
                new_doc.insert_pdf(doc, from_page=data['original_num']-1, to_page=data['original_num']-1)
            
            new_doc.save(output_path)
            new_doc.close()
            doc.close()
            
            return output_path
            
        except Exception as e:
            raise Exception(f"ಪುಟ ಸಾರಿಸುವಿಕೆ ವಿಫಲ: {str(e)}")

    def get_page_sorting_preview(self, file_path, session_id):
        """Generate preview for page sorting"""
        try:
            from .kannada_numeral_converter import KannadaNumeralConverter
            
            converter = KannadaNumeralConverter()
            doc = fitz.open(file_path)
            
            previews = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                extracted_number = converter.extract_page_number_from_text(text)
                
                thumbnail_path = self._generate_page_thumbnail(page, page_num + 1, session_id)
                
                previews.append({
                    'page_num': page_num + 1,
                    'extracted_num': extracted_number,
                    'thumbnail_path': thumbnail_path
                })
            
            doc.close()
            
            return {
                'total_pages': len(previews),
                'previews': previews
            }
            
        except Exception as e:
            return {'error': f'ಪೂರ್ವವೀಕ್ಷಣೆ ರಚನೆ ವಿಫಲ: {str(e)}'}

    def _generate_page_thumbnail(self, page, page_num, session_id):
        """Generate a thumbnail image for a PDF page"""
        try:
            import os
            from PIL import Image
            import io
            
            thumbnails_dir = os.path.join(self.config.OUTPUT_FOLDER, 'thumbnails', session_id)
            os.makedirs(thumbnails_dir, exist_ok=True)
            
            import time
            timestamp = int(time.time() * 1000)
            thumbnail_filename = f"page_{page_num}_{timestamp}.png"
            thumbnail_path = os.path.join(thumbnails_dir, thumbnail_filename)
            
            zoom = 1.5
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # Apply 180-degree rotation to fix upside-down orientation
            img = img.rotate(180, expand=True)
            
            thumbnail_size = (150, 200)
            img.thumbnail(thumbnail_size, Image.Resampling.LANCZOS)
            
            img.save(thumbnail_path, "PNG", optimize=True)
            
            return f'/thumbnails/{session_id}/{thumbnail_filename}'
            
        except Exception as e:
            return None

    def add_watermark(self, file_path, session_id, watermark_options):
        """Add watermark to PDF"""
        try:
            from .validators import validate_watermark_options
            
            is_valid, message = validate_watermark_options(watermark_options)
            if not is_valid:
                raise Exception(message)
            
            output_filename = f"{session_id}_watermarked.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                if watermark_options['type'] == 'text':
                    self._add_text_watermark(page, watermark_options)
                else:
                    self._add_image_watermark(page, watermark_options)
            
            doc.save(output_path)
            doc.close()
            
            return output_path
            
        except Exception as e:
            raise Exception(f"ವಾಟರ್‌ಮಾರ್ಕ್ ಸೇರಿಸುವಿಕೆ ವಿಫಲ: {str(e)}")

    def _add_text_watermark(self, page, options):
        """Add text watermark to page"""
        rect = page.rect
        
        if options['position'] == 'center':
            x, y = rect.width / 2, rect.height / 2
        elif options['position'] == 'top-left':
            x, y = 50, rect.height - 50
        elif options['position'] == 'top-right':
            x, y = rect.width - 50, rect.height - 50
        elif options['position'] == 'bottom-left':
            x, y = 50, 50
        elif options['position'] == 'bottom-right':
            x, y = rect.width - 50, 50
        else:
            x, y = rect.width / 2, rect.height / 2
        
        page.insert_text(
            (x, y),
            options['text'],
            fontsize=options['font_size'],
            color=self._hex_to_rgb(options['color']),
            rotate=options['rotation']
        )

    def _add_image_watermark(self, page, options):
        """Add image watermark to page"""
        if 'image_path' not in options or not os.path.exists(options['image_path']):
            return
        
        rect = page.rect
        
        if options['position'] == 'center':
            x, y = rect.width / 2 - 50, rect.height / 2 - 50
        else:
            x, y = rect.width / 2 - 50, rect.height / 2 - 50
        
        image_rect = fitz.Rect(x, y, x + 100, y + 100)
        page.insert_image(image_rect, filename=options['image_path'])

    def _hex_to_rgb(self, hex_color):
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16)/255.0 for i in (0, 2, 4))
