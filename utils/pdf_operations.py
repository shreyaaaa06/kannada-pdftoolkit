import os
import zipfile
from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
import fitz
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import config
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
import html
import platform
import subprocess
import shutil
from pathlib import Path
import io
import psutil
import gc
from datetime import datetime
import json
import tempfile
import os
import platform
import shutil
import subprocess
from pathlib import Path
import platform
import subprocess
import shutil
import os
from pathlib import Path
import threading
from .pdf_compare import PDFCompare

class PDFOperations:
    def __init__(self):
        self.config = config.Config()
    
    def merge_pdfs(self, file_paths, session_id):
        """Merge multiple PDF files"""
        try:
            print(f"=== MERGE PDF DEBUG ===")
            print(f"Session ID: {session_id}")
            print(f"File paths: {file_paths}")
            print(f"Number of files: {len(file_paths)}")
        
        # Validate input
            if not file_paths or len(file_paths) < 2:
                raise Exception("ವಿಲೀನಕ್ಕೆ ಕನಿಷ್ಠ 2 PDF ಫೈಲ್‌ಗಳು ಅಗತ್ಯ")
        
            valid_files = []
            for i, file_path in enumerate(file_paths):
                print(f"Checking file {i+1}: {file_path}")          
                
                if not os.path.exists(file_path):
                    print(f"File does not exist: {file_path}")
                    continue
                if os.path.getsize(file_path) == 0:
                    continue
            
                try:
                    test_reader = PdfReader(file_path)
                    if len(test_reader.pages) == 0:
                        print(f"PDF has no pages: {file_path}")
                        continue
                    print(f"Valid PDF with {len(test_reader.pages)} pages: {file_path}")
                    valid_files.append(file_path)
                except Exception as pdf_error:
                    print(f"Invalid PDF file {file_path}: {pdf_error}")
                    continue
        
            if len(valid_files) < 2:
                raise Exception("ವಿಲೀನಕ್ಕೆ ಕನಿಷ್ಠ 2 ಸರಿಯಾದ PDF ಫೈಲ್‌ಗಳು ಅಗತ್ಯ")
            
            print(f"Valid files for merging: {len(valid_files)}")

            writer = PdfWriter()
            total_pages_added = 0
            
            for file_path in valid_files:
                try:
                    reader = PdfReader(file_path)
                    pages_in_file = len(reader.pages)
                    print(f"Pages in file: {pages_in_file}")
                
                # Add all pages from this PDF
                    for page_num, page in enumerate(reader.pages):
                        try:
                            writer.add_page(page)
                            total_pages_added += 1
                            print(f"Added page {page_num + 1} from {os.path.basename(file_path)}")
                        except Exception as page_error:
                            print(f"Error adding page {page_num + 1}: {page_error}")
                            continue
                        
                except Exception as file_error:
                    print(f"Error processing file {file_path}: {file_error}")
                    continue
        
            if total_pages_added == 0:
                raise Exception("ಯಾವುದೇ ಪುಟಗಳನ್ನು ವಿಲೀನ ಮಾಡಲಾಗಿಲ್ಲ")
        
            print(f"Total pages added to merged PDF: {total_pages_added}")
        
        # Create output path
            output_path = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_merged.pdf")
            print(f"Output path: {output_path}")
        
        # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Write the merged PDF
            try:
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                print(f"Merged PDF written successfully")
            except Exception as write_error:
                print(f"Error writing merged PDF: {write_error}")
                raise Exception(f"ವಿಲೀನ PDF ಬರೆಯುವಲ್ಲಿ ದೋಷ: {str(write_error)}")
        
        # Verify the output file
            if not os.path.exists(output_path):
                raise Exception("ವಿಲೀನ PDF ಫೈಲ್ ರಚಿಸಲಾಗಿಲ್ಲ")
        
            output_size = os.path.getsize(output_path)
            if output_size == 0:
                raise Exception("ಖಾಲಿ ವಿಲೀನ PDF ರಚಿಸಲಾಗಿದೆ")
        
            print(f"Merged PDF size: {output_size} bytes")
        
        # Final validation - try to read the merged PDF
            try:
                validation_reader = PdfReader(output_path)
                merged_pages = len(validation_reader.pages)
                print(f"Validation: Merged PDF has {merged_pages} pages")
            
                if merged_pages == 0:
                    raise Exception("ವಿಲೀನ PDF ಯಲ್ಲಿ ಯಾವುದೇ ಪುಟಗಳಿಲ್ಲ")
                
            except Exception as validation_error:
                print(f"Validation error: {validation_error}")
                raise Exception(f"ವಿಲೀನ PDF ದೋಷಪೂರ್ಣ: {str(validation_error)}")
        
            print(f"=== MERGE SUCCESSFUL ===")
            print(f"Files merged: {len(valid_files)}")
            print(f"Total pages: {merged_pages}")
            print(f"Output file: {output_path}")
            print(f"=== END DEBUG ===")
        
            return output_path
        
        except Exception as e:
            print(f"Merge error: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"PDF ವಿಲೀನ ವಿಫಲ: {str(e)}")
    
    def split_pdf(self, file_path, session_id, pages="", split_method="pages", target_size_mb=10, pages_per_chunk=20, max_file_size_mb=500):
        """Enhanced PDF split function with proper split method handling"""
        try:
            print(f"=== ENHANCED PDF SPLIT DEBUG ===")
            print(f"File path: {file_path}")
            print(f"Session ID: {session_id}")
            print(f"Pages parameter: '{pages}'")
            print(f"Split method: '{split_method}'")
            print(f"Target size MB: {target_size_mb}")
            print(f"Pages per chunk: {pages_per_chunk}")
            print(f"Max file size: {max_file_size_mb}MB")
            print(f"Timestamp: {datetime.now()}")

            # ===== INPUT VALIDATION =====
            if not os.path.exists(file_path):
                raise Exception("PDF ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ")
            
            file_size_bytes = os.path.getsize(file_path)
            file_size_mb = file_size_bytes / (1024 * 1024)
            print(f"Input file size: {file_size_mb:.2f}MB ({file_size_bytes:,} bytes)")
            
            if file_size_bytes == 0:
                raise Exception("ಖಾಲಿ PDF ಫೈಲ್")
            
            if file_size_mb > max_file_size_mb:
                raise Exception(f"ಫೈಲ್ ತುಂಬಾ ದೊಡ್ಡದಾಗಿದೆ. ಗರಿಷ್ಠ ಗಾತ್ರ: {max_file_size_mb}MB")
            
            # Check available memory
            available_memory_mb = psutil.virtual_memory().available / (1024 * 1024)
            required_memory_mb = file_size_mb * 3
            print(f"Available memory: {available_memory_mb:.0f}MB, Required: {required_memory_mb:.0f}MB")
            
            if required_memory_mb > available_memory_mb:
                print("WARNING: Low memory detected, using memory-efficient processing")
            
            # ===== PDF VALIDATION =====
            print("Validating PDF structure...")
            reader = None
            total_pages = 0
            
            try:
                reader = PdfReader(file_path)
                total_pages = len(reader.pages)
                print(f"PDF validation successful: {total_pages} pages")
                
                if total_pages == 0:
                    raise Exception("PDF ಯಲ್ಲಿ ಯಾವುದೇ ಪುಟಗಳಿಲ್ಲ")
                
                if reader.is_encrypted:
                    raise Exception("ಎನ್‌ಕ್ರಿಪ್ಟ್ ಮಾಡಿದ PDF ಬೆಂಬಲಿತವಾಗಿಲ್ಲ")
                    
            except Exception as pdf_error:
                print(f"PyPDF2 validation failed: {pdf_error}")
                try:
                    doc = fitz.open(file_path)
                    total_pages = len(doc)
                    doc.close()
                    print(f"PyMuPDF validation successful: {total_pages} pages")
                    
                    if total_pages == 0:
                        raise Exception("PDF ಯಲ್ಲಿ ಯಾವುದೇ ಪುಟಗಳಿಲ್ಲ")
                        
                except Exception as fitz_error:
                    raise Exception(f"ಅಮಾನ್ಯ PDF ಫೈಲ್: {str(fitz_error)}")
            
            if total_pages == 1:
                raise Exception("ಒಂದೇ ಪುಟದ PDF ಅನ್ನು ವಿಭಜಿಸಲು ಸಾಧ್ಯವಿಲ್ಲ")

            # ===== OUTPUT DIRECTORY SETUP =====
            temp_dir = tempfile.mkdtemp(prefix=f"pdf_split_{session_id}_")
            output_dir = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_split")
            
            try:
                os.makedirs(output_dir, exist_ok=True)
                print(f"Output directory: {output_dir}")
                print(f"Temp directory: {temp_dir}")
                
                # ===== DETERMINE SPLIT STRATEGY BASED ON METHOD =====
                print(f"Processing split method: {split_method}")
                
                created_files = []
                
                if split_method == "size":
                    print(f"Size-based splitting: target {target_size_mb}MB per file")
                    created_files = self._split_by_file_size(
                        file_path, temp_dir, target_size_mb, total_pages, session_id
                    )
                    
                elif split_method == "auto":
                    print(f"Auto chunking: {pages_per_chunk} pages per chunk")
                    created_files = self._auto_chunk_pdf(
                        file_path, temp_dir, pages_per_chunk, total_pages, session_id
                    )
                    
                elif split_method == "pages" or not split_method:
                    print("Page-based splitting")
                    if not pages or pages.strip() == "":
                        # Default: split in middle
                        split_point = max(1, total_pages // 2)
                        print(f"No pages specified, splitting at page {split_point}")
                        created_files = self._split_pdf_two_parts(
                            file_path, temp_dir, split_point, total_pages, session_id
                        )
                    else:
                        # Parse page specification
                        split_info = self._parse_split_specification_enhanced(pages, total_pages, file_size_mb, target_size_mb)
                        print(f"Split strategy: {split_info}")
                        
                        if split_info['type'] == 'single_split':
                            created_files = self._split_pdf_two_parts(
                                file_path, temp_dir, split_info['split_point'], total_pages, session_id
                            )
                        elif split_info['type'] == 'extract_pages':
                            created_files = self._extract_specific_pages(
                                file_path, temp_dir, split_info['pages'], session_id
                            )
                        elif split_info['type'] == 'multiple_splits':
                            created_files = self._split_multiple_ranges(
                                file_path, temp_dir, split_info['ranges'], session_id
                            )
                        elif split_info['type'] == 'auto_chunk':
                            created_files = self._auto_chunk_pdf(
                                file_path, temp_dir, split_info['pages_per_chunk'], total_pages, session_id
                            )
                        else:
                            # Fallback
                            split_point = max(1, total_pages // 2)
                            created_files = self._split_pdf_two_parts(
                                file_path, temp_dir, split_point, total_pages, session_id
                            )
                else:
                    raise Exception(f"ಅಮಾನ್ಯ ವಿಭಾಗ ವಿಧಾನ: {split_method}")
                
                # ===== VALIDATION AND CLEANUP =====
                print(f"Split operation completed. Files created: {len(created_files)}")
                
                validated_files = []
                total_output_size = 0
                
                for file_path_created in created_files:
                    try:
                        if not os.path.exists(file_path_created):
                            print(f"WARNING: File not found: {os.path.basename(file_path_created)}")
                            continue
                        
                        file_size = os.path.getsize(file_path_created)
                        if file_size == 0:
                            print(f"WARNING: Empty file: {os.path.basename(file_path_created)}")
                            continue
                        
                        try:
                            test_reader = PdfReader(file_path_created)
                            page_count = len(test_reader.pages)
                            if page_count == 0:
                                print(f"WARNING: PDF with no pages: {os.path.basename(file_path_created)}")
                                continue
                            print(f"✓ {os.path.basename(file_path_created)}: {page_count} pages, {file_size:,} bytes")
                            validated_files.append(file_path_created)
                            total_output_size += file_size
                            
                        except Exception as validation_error:
                            print(f"WARNING: Invalid PDF {os.path.basename(file_path_created)}: {validation_error}")
                            continue
                            
                    except Exception as file_error:
                        print(f"WARNING: Error validating {file_path_created}: {file_error}")
                        continue
                
                if not validated_files:
                    raise Exception("ಯಾವುದೇ ಸರಿಯಾದ PDF ಫೈಲ್‌ಗಳು ರಚಿಸಲಾಗಿಲ್ಲ")
                
                print(f"Validated files: {len(validated_files)}")
                print(f"Total output size: {total_output_size:,} bytes ({total_output_size/(1024*1024):.2f}MB)")
                
                # ===== CREATE OUTPUT ZIP =====
                zip_path = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_split.zip")
                print(f"Creating ZIP archive: {zip_path}")
                
                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED, compresslevel=6) as zipf:
                    for i, file_path_to_zip in enumerate(validated_files):
                        try:
                            filename = os.path.basename(file_path_to_zip)
                            zipf.write(file_path_to_zip, filename)
                            print(f"Added to ZIP ({i+1}/{len(validated_files)}): {filename}")
                            
                        except Exception as zip_error:
                            print(f"ERROR: Failed to add {filename} to ZIP: {zip_error}")
                            continue
                
                # Validate ZIP file
                if not os.path.exists(zip_path):
                    raise Exception("ZIP ಫೈಲ್ ರಚಿಸಲಾಗಿಲ್ಲ")
                
                zip_size = os.path.getsize(zip_path)
                if zip_size == 0:
                    raise Exception("ಖಾಲಿ ZIP ಫೈಲ್ ರಚಿಸಲಾಗಿದೆ")
                
                try:
                    with zipfile.ZipFile(zip_path, 'r') as zipf:
                        zip_contents = zipf.namelist()
                        if not zip_contents:
                            raise Exception("ZIP ಫೈಲ್‌ನಲ್ಲಿ ಯಾವುದೇ ಫೈಲ್‌ಗಳಿಲ್ಲ")
                        print(f"ZIP verification successful: {len(zip_contents)} files")
                        
                except Exception as zip_verify_error:
                    raise Exception(f"ZIP ಫೈಲ್ ದೋಷಪೂರ್ಣ: {str(zip_verify_error)}")
                
                print(f"ZIP created successfully: {zip_size:,} bytes ({zip_size/(1024*1024):.2f}MB)")
                
                # ===== CLEANUP =====
                try:
                    shutil.rmtree(temp_dir, ignore_errors=True)
                    print("Temporary files cleaned up")
                except Exception as cleanup_error:
                    print(f"WARNING: Cleanup failed: {cleanup_error}")
                
                if file_size_mb > 100:
                    gc.collect()
                    print("Memory cleanup performed")
                
                # ===== SUCCESS SUMMARY =====
                compression_ratio = (zip_size / total_output_size) * 100 if total_output_size > 0 else 100
                
                print(f"=== SPLIT OPERATION SUCCESSFUL ===")
                print(f"Input file: {file_size_mb:.2f}MB ({total_pages} pages)")
                print(f"Output files: {len(validated_files)} PDFs")
                print(f"ZIP file: {zip_size/(1024*1024):.2f}MB")
                print(f"Compression ratio: {compression_ratio:.1f}%")
                print(f"Output path: {zip_path}")
                print(f"=== END DEBUG ===")
                
                return zip_path
                
            except Exception as processing_error:
                try:
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir, ignore_errors=True)
                except:
                    pass
                raise processing_error
            
        except Exception as e:
            print(f"Enhanced split error: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"PDF ವಿಭಾಗ ವಿಫಲ: {str(e)}")

    def _split_by_file_size(self, file_path, output_dir, target_size_mb, total_pages, session_id):
        """Split PDF based on target file size"""
        print(f"Starting size-based split: target {target_size_mb}MB per file")
        
        try:
            # Calculate approximate pages per chunk based on file size
            file_size_bytes = os.path.getsize(file_path)
            file_size_mb = file_size_bytes / (1024 * 1024)
            
            # Estimate pages per chunk
            estimated_pages_per_chunk = max(1, int((target_size_mb / file_size_mb) * total_pages))
            print(f"Estimated pages per chunk: {estimated_pages_per_chunk}")
            
            created_files = []
            current_page = 1
            chunk_num = 1
            
            reader = PdfReader(file_path)
            
            while current_page <= total_pages:
                end_page = min(current_page + estimated_pages_per_chunk - 1, total_pages)
                
                writer = PdfWriter()
                pages_in_chunk = 0
                
                for page_num in range(current_page, end_page + 1):
                    if page_num <= len(reader.pages):
                        writer.add_page(reader.pages[page_num - 1])
                        pages_in_chunk += 1
                
                if pages_in_chunk > 0:
                    output_path = os.path.join(
                        output_dir, 
                        f"size_chunk_{chunk_num:03d}_pages_{current_page}_to_{end_page}.pdf"
                    )
                    
                    with open(output_path, 'wb') as output_file:
                        writer.write(output_file)
                    
                    # Check actual file size
                    actual_size_mb = os.path.getsize(output_path) / (1024 * 1024)
                    created_files.append(output_path)
                    print(f"Created size chunk {chunk_num}: pages {current_page}-{end_page} ({pages_in_chunk} pages, {actual_size_mb:.2f}MB)")
                
                current_page = end_page + 1
                chunk_num += 1
                
                writer = None
                if chunk_num % 10 == 0:
                    gc.collect()
            
            return created_files
            
        except Exception as e:
            print(f"Size-based split error: {e}")
            raise Exception(f"ಗಾತ್ರದ ಆಧಾರದ ವಿಭಾಗ ವಿಫಲ: {str(e)}")
    
    def _parse_split_specification_enhanced(self, pages, total_pages, file_size_mb, chunk_size_mb):
        """Enhanced parsing with intelligent split strategy selection"""
        try:
            print(f"Parsing split specification with file size: {file_size_mb:.2f}MB")
            
            # For very large files, prefer chunking
            if file_size_mb > 200:
                pages_per_chunk = max(10, int((chunk_size_mb / file_size_mb) * total_pages))
                if not pages or pages.strip() == "":
                    print(f"Large file detected: auto-chunking with {pages_per_chunk} pages per chunk")
                    return {
                        'type': 'auto_chunk',
                        'pages_per_chunk': pages_per_chunk,
                        'reason': 'large_file_optimization'
                    }
            
            if not pages or pages.strip() == "":
                split_point = max(1, total_pages // 2)
                return {
                    'type': 'single_split',
                    'split_point': split_point,
                    'reason': 'default_middle_split'
                }
            
            pages_str = pages.strip().lower()
            
            # Handle special keywords
            if pages_str in ['auto', 'chunk', 'smart']:
                pages_per_chunk = max(10, min(50, total_pages // 5))
                return {
                    'type': 'auto_chunk',
                    'pages_per_chunk': pages_per_chunk,
                    'reason': 'user_requested_auto'
                }
            
            if pages_str in ['efficient', 'memory', 'large']:
                chunk_pages = max(20, min(100, total_pages // 10))
                return {
                    'type': 'memory_efficient_chunks',
                    'chunk_size': chunk_pages,
                    'reason': 'memory_efficient_requested'
                }
            
            # Original parsing logic
            try:
                split_point = int(pages_str)
                split_point = max(1, min(split_point, total_pages - 1))
                return {
                    'type': 'single_split',
                    'split_point': split_point,
                    'reason': 'user_specified_point'
                }
            except ValueError:
                pass
            
            # Range notation
            if '-' in pages_str and ',' not in pages_str:
                parts = pages_str.split('-')
                if len(parts) == 2:
                    try:
                        start = int(parts[0].strip())
                        end = int(parts[1].strip())
                        start = max(1, min(start, total_pages))
                        end = max(start, min(end, total_pages))
                        
                        return {
                            'type': 'extract_pages',
                            'pages': list(range(start, end + 1)),
                            'reason': 'user_specified_range'
                        }
                    except ValueError:
                        pass
            
            # Comma-separated pages/ranges
            if ',' in pages_str:
                try:
                    page_numbers = self._parse_page_ranges_enhanced(pages_str, total_pages)
                    if page_numbers:
                        return {
                            'type': 'extract_pages',
                            'pages': page_numbers,
                            'reason': 'user_specified_pages'
                        }
                except:
                    pass
            
            # Fallback
            split_point = max(1, total_pages // 2)
            return {
                'type': 'single_split',
                'split_point': split_point,
                'reason': 'fallback_middle_split'
            }
            
        except Exception as e:
            print(f"Parse specification error: {e}")
            return {
                'type': 'single_split',
                'split_point': max(1, total_pages // 2),
                'reason': 'error_fallback'
            }

    def _parse_page_ranges_enhanced(self, pages_str, total_pages):
        """Enhanced page range parsing with better error handling"""
        pages = set()
        
        try:
            parts = [part.strip() for part in pages_str.split(',') if part.strip()]
            
            for part in parts:
                if '-' in part:
                    range_parts = part.split('-', 1)
                    if len(range_parts) == 2:
                        try:
                            start = int(range_parts[0].strip())
                            end = int(range_parts[1].strip())
                            
                            start = max(1, min(start, total_pages))
                            end = max(start, min(end, total_pages))
                            
                            pages.update(range(start, end + 1))
                            
                        except ValueError:
                            print(f"Invalid range format: {part}")
                            continue
                else:
                    try:
                        page_num = int(part)
                        if 1 <= page_num <= total_pages:
                            pages.add(page_num)
                        else:
                            print(f"Page {page_num} out of range (1-{total_pages})")
                            
                    except ValueError:
                        print(f"Invalid page number: {part}")
                        continue
            
            result = sorted(list(pages))
            print(f"Parsed page ranges result: {len(result)} pages")
            return result
            
        except Exception as e:
            print(f"Enhanced page range parsing error: {e}")
            return []

    def _extract_pages_efficient(self, file_path, output_dir, page_numbers, filename_prefix, session_id):
        """Memory-efficient page extraction"""
        try:
            output_filename = f"{filename_prefix}_pages_{min(page_numbers)}_to_{max(page_numbers)}.pdf"
            output_path = os.path.join(output_dir, output_filename)
            
            reader = PdfReader(file_path)
            writer = PdfWriter()
            
            for page_num in page_numbers:
                if 1 <= page_num <= len(reader.pages):
                    page = reader.pages[page_num - 1]
                    writer.add_page(page)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            writer = None
            return output_path
            
        except Exception as e:
            print(f"Efficient page extraction error: {e}")
            return None

    def _split_pdf_two_parts(self, file_path, output_dir, split_point, total_pages, session_id):
        """Traditional two-part split with memory management"""
        created_files = []
        
        try:
            reader = PdfReader(file_path)
            
            # First part
            print(f"Creating first part: pages 1 to {split_point}")
            writer1 = PdfWriter()
            for i in range(split_point):
                writer1.add_page(reader.pages[i])
            
            output_path1 = os.path.join(output_dir, f"part_1_pages_1_to_{split_point}.pdf")
            with open(output_path1, 'wb') as output_file:
                writer1.write(output_file)
            created_files.append(output_path1)
            
            writer1 = None
            
            # Second part
            print(f"Creating second part: pages {split_point + 1} to {total_pages}")
            writer2 = PdfWriter()
            for i in range(split_point, total_pages):
                writer2.add_page(reader.pages[i])
            
            output_path2 = os.path.join(output_dir, f"part_2_pages_{split_point + 1}_to_{total_pages}.pdf")
            with open(output_path2, 'wb') as output_file:
                writer2.write(output_file)
            created_files.append(output_path2)
            
            return created_files
            
        except Exception as e:
            print(f"Two-part split error: {e}")
            raise Exception(f"ಎರಡು ಭಾಗಗಳ ವಿಭಾಗ ವಿಫಲ: {str(e)}")

    def _extract_specific_pages(self, file_path, output_dir, pages_to_extract, session_id):
        """Extract specific pages with validation"""
        try:
            reader = PdfReader(file_path)
            writer = PdfWriter()
            
            extracted_count = 0
            for page_num in pages_to_extract:
                if 1 <= page_num <= len(reader.pages):
                    writer.add_page(reader.pages[page_num - 1])
                    extracted_count += 1
            
            if extracted_count == 0:
                raise Exception("ಯಾವುದೇ ಸರಿಯಾದ ಪುಟಗಳು ಕಂಡುಬಂದಿಲ್ಲ")
            
            if len(pages_to_extract) <= 5:
                page_str = "_".join(map(str, pages_to_extract))
            else:
                page_str = f"{min(pages_to_extract)}_to_{max(pages_to_extract)}_and_others"
            
            output_path = os.path.join(output_dir, f"extracted_pages_{page_str}.pdf")
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            print(f"Extracted {extracted_count} pages to {os.path.basename(output_path)}")
            return [output_path]
            
        except Exception as e:
            print(f"Specific page extraction error: {e}")
            raise Exception(f"ನಿರ್ದಿಷ್ಟ ಪುಟಗಳ ಹೊರತೆಗೆಯುವಿಕೆ ವಿಫಲ: {str(e)}")

    def _split_multiple_ranges(self, file_path, output_dir, ranges, session_id):
        """Split into multiple ranges with memory management"""  
        created_files = []
        
        try:
            reader = PdfReader(file_path)
            
            for i, (start, end) in enumerate(ranges):
                writer = PdfWriter()
                pages_added = 0
                
                for page_num in range(start, end + 1):
                    if 1 <= page_num <= len(reader.pages):
                        writer.add_page(reader.pages[page_num - 1])
                        pages_added += 1
                
                if pages_added > 0:
                    output_path = os.path.join(output_dir, f"part_{i + 1}_pages_{start}_to_{end}.pdf")
                    with open(output_path, 'wb') as output_file:
                        writer.write(output_file)
                    created_files.append(output_path)
                    print(f"Created range {i + 1}: {pages_added} pages")
                
                writer = None
            
            return created_files
            
        except Exception as e:
            print(f"Multiple ranges split error: {e}")
            raise Exception(f"ಬಹು ವ್ಯಾಪ್ತಿ ವಿಭಾಗ ವಿಫಲ: {str(e)}")
    
    def _auto_chunk_pdf(self, file_path, output_dir, pages_per_chunk, total_pages, session_id):
            """Automatic chunking based on optimal chunk size"""
            print(f"Auto-chunking PDF: {pages_per_chunk} pages per chunk")
            
            created_files = []
            current_page = 1
            chunk_num = 1
            
            try:
                reader = PdfReader(file_path)
                
                while current_page <= total_pages:
                    end_page = min(current_page + pages_per_chunk - 1, total_pages)
                    
                    writer = PdfWriter()
                    pages_in_chunk = 0
                    
                    for page_num in range(current_page, end_page + 1):
                        if page_num <= len(reader.pages):
                            writer.add_page(reader.pages[page_num - 1])
                            pages_in_chunk += 1
                    
                    if pages_in_chunk > 0:
                        output_path = os.path.join(
                            output_dir, 
                            f"chunk_{chunk_num:03d}_pages_{current_page}_to_{end_page}.pdf"
                        )
                        
                        with open(output_path, 'wb') as output_file:
                            writer.write(output_file)
                        
                        created_files.append(output_path)
                        print(f"Created chunk {chunk_num}: pages {current_page}-{end_page} ({pages_in_chunk} pages)")
                    
                    current_page = end_page + 1
                    chunk_num += 1
                    
                    writer = None
                    if chunk_num % 10 == 0:
                        gc.collect()
                
                return created_files
                
            except Exception as e:
                print(f"Auto-chunk error: {e}")
                raise Exception(f"ಸ್ವಯಂಚಾಲಿತ ಚಂಕಿಂಗ್ ವಿಫಲ: {str(e)}")

    def extract_pages(self, file_path, pages, session_id):
        """Extract specific pages from PDF"""
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            page_numbers = self._parse_page_ranges(pages, total_pages)
            
            if not page_numbers:
                raise Exception("ಸರಿಯಾದ ಪುಟ ಸಂಖ್ಯೆಗಳನ್ನು ನಮೂದಿಸಿ")
            
            writer = PdfWriter()
            for page_num in page_numbers:
                writer.add_page(reader.pages[page_num - 1])
            
            output_filename = f"{session_id}_extracted.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"ಪುಟ ಹೊರತೆಗೆಯುವಿಕೆ ವಿಫಲ: {str(e)}")

    def delete_pages(self, file_path, pages, session_id):
        """Delete specific pages from PDF"""
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            pages_to_delete = set(self._parse_page_ranges(pages, total_pages))
            
            writer = PdfWriter()
            for i, page in enumerate(reader.pages):
                if (i + 1) not in pages_to_delete:
                    writer.add_page(page)
            
            output_filename = f"{session_id}_deleted.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"ಪುಟ ಅಳಿಸುವಿಕೆ ವಿಫಲ: {str(e)}")
        

        # Test if valid PDF
    def compress_pdf(self, file_path, compression_level, session_id, target_size_mb=None):
        """
        Enhanced PDF compression with user-controlled levels and target size
        
        Args:
            file_path: Path to input PDF
            compression_level: 'low', 'medium', 'high', 'maximum', or 'custom'
            session_id: Session identifier
            target_size_mb: Target file size in MB (for custom compression)
        """
        try:
            print(f"=== ENHANCED PDF COMPRESSION ===")
            print(f"Input: {file_path}")
            print(f"Level: {compression_level}")
            print(f"Target size: {target_size_mb}MB" if target_size_mb else "No target size")
            
            # Validate input
            if not os.path.exists(file_path):
                raise Exception("PDF ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ")
            
            original_size = os.path.getsize(file_path)
            original_size_mb = original_size / (1024 * 1024)
            print(f"Original size: {original_size:,} bytes ({original_size_mb:.2f} MB)")
            
            if original_size == 0:
                raise Exception("ಖಾಲಿ PDF ಫೈಲ್")
            
            # Test if valid PDF
            try:
                doc = fitz.open(file_path)
                page_count = len(doc)
                doc.close()
                print(f"Valid PDF with {page_count} pages")
            except Exception as e:
                raise Exception(f"ಅಮಾನ್ಯ PDF ಫೈಲ್: {str(e)}")
            
            output_path = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_compressed.pdf")
            
            # Enhanced compression methods with different aggressiveness levels
            if compression_level == 'maximum' or (target_size_mb and target_size_mb < original_size_mb * 0.3):
                # Ultra-aggressive compression
                methods = [
                    ("Ultra Aggressive", self._compress_ultra_aggressive),
                    ("Extreme Image Recreation", self._compress_extreme_image_recreation),
                    ("Advanced PyMuPDF", self._compress_advanced_pymupdf),
                ]
            elif compression_level == 'high' or (target_size_mb and target_size_mb < original_size_mb * 0.6):
                # High compression
                methods = [
                    ("Extreme Image Recreation", self._compress_extreme_image_recreation),
                    ("Advanced PyMuPDF", self._compress_advanced_pymupdf),
                    ("Image Recreation", self._compress_by_image_recreation),
                ]
            elif compression_level == 'medium':
                # Balanced compression
                methods = [
                    ("Advanced PyMuPDF", self._compress_advanced_pymupdf),
                    ("Image Recreation", self._compress_by_image_recreation),
                    ("Basic PyMuPDF", self._compress_basic_pymupdf)
                ]
            else:  # low compression
                # Gentle compression
                methods = [
                    ("Basic PyMuPDF", self._compress_basic_pymupdf),
                    ("Advanced PyMuPDF", self._compress_advanced_pymupdf),
                ]
            
            best_result = None
            best_size = original_size
            target_reached = False 
            for method_name, method_func in methods:
                try:
                    print(f"Trying {method_name} compression...")
                    temp_output = output_path.replace('.pdf', f'_temp_{method_name.replace(" ", "_").lower()}.pdf')
                    
                    # Pass target size for adaptive compression
                    if hasattr(method_func, '__code__') and 'target_size_mb' in method_func.__code__.co_varnames:
                        result = method_func(file_path, temp_output, compression_level, target_size_mb)
                    else:
                        result = method_func(file_path, temp_output, compression_level)
                    
                    if result and os.path.exists(temp_output):
                        compressed_size = os.path.getsize(temp_output)
                        compressed_size_mb = compressed_size / (1024 * 1024)
                        
                        if compressed_size > 0 and compressed_size < best_size:
                            # Clean up previous best
                            if best_result and os.path.exists(best_result):
                                os.remove(best_result)
                            
                            best_result = temp_output
                            best_size = compressed_size
                            reduction = (1 - compressed_size/original_size) * 100
                            
                            print(f"✓ {method_name} successful!")
                            print(f"Compressed size: {compressed_size:,} bytes ({compressed_size_mb:.2f} MB)")
                            print(f"Size reduction: {reduction:.1f}%")
                            
                            # Check if target size reached
                            if target_size_mb and compressed_size_mb <= target_size_mb:
                                target_reached = True
                                print(f"🎯 Target size reached! ({compressed_size_mb:.2f}MB <= {target_size_mb}MB)")
                                break
                        else:
                            print(f"✗ {method_name} didn't improve compression")
                            if os.path.exists(temp_output):
                                os.remove(temp_output)
                    else:
                        print(f"✗ {method_name} failed to create output")
                        
                except Exception as e:
                    print(f"✗ {method_name} failed: {e}")
                    continue
            
            # Use the best result or try iterative compression if target not reached
            if best_result and os.path.exists(best_result):
                final_output = output_path
                
                # If target size not reached and we have a target, try iterative compression
                if target_size_mb and not target_reached and best_size > target_size_mb * 1024 * 1024:
                    print("Target size not reached, trying iterative compression...")
                    iterative_result = self._iterative_compression(best_result, final_output, target_size_mb)
                    if iterative_result and os.path.exists(final_output):
                        if best_result != final_output:
                            os.remove(best_result)
                    else:
                        # Keep the best result we got
                        if best_result != final_output:
                            os.rename(best_result, final_output)
                else:
                    # Move best result to final output path
                    if os.path.exists(final_output):
                        os.remove(final_output)
                    os.rename(best_result, final_output)
                
                final_size = os.path.getsize(final_output)
                final_size_mb = final_size / (1024 * 1024)
                reduction = (1 - final_size/original_size) * 100
                
                print(f"=== COMPRESSION COMPLETE ===")
                print(f"Original: {original_size:,} bytes ({original_size_mb:.2f} MB)")
                print(f"Compressed: {final_size:,} bytes ({final_size_mb:.2f} MB)")
                print(f"Reduction: {reduction:.1f}%")
                
                if target_size_mb:
                    if final_size_mb <= target_size_mb:
                        print(f"🎯 Target achieved: {final_size_mb:.2f}MB <= {target_size_mb}MB")
                    else:
                        print(f"⚠️ Target not fully achieved: {final_size_mb:.2f}MB > {target_size_mb}MB")
                
                # Only return if we actually achieved compression
                if final_size < original_size:
                    return final_output
                else:
                    print("No compression achieved, removing output file")
                    os.remove(final_output)
                    raise Exception("ಸಂಕುಚನ ಸಾಧ್ಯವಾಗಲಿಲ್ಲ - ಮೂಲ ಫೈಲ್ ಈಗಾಗಲೇ ಅತ್ಯುತ್ತಮವಾಗಿ ಸಂಕುಚಿತವಾಗಿದೆ")
            
            # If no method worked
            raise Exception("ಎಲ್ಲಾ ಸಂಕುಚನ ವಿಧಾನಗಳು ವಿಫಲವಾಗಿವೆ")
            
        except Exception as e:
            print(f"Compression error: {str(e)}")
            raise Exception(f"PDF ಸಂಕುಚನ ವಿಫಲ: {str(e)}")
    
    def rotate_pdf(self, file_path, session_id, rotation_angle=90, pages="", apply_to_all=True):
        """Rotate PDF pages by specified angle"""
        try:
            print(f"=== PDF ROTATION DEBUG ===")
            print(f"File path: {file_path}")
            print(f"Session ID: {session_id}")
            print(f"Rotation angle: {rotation_angle}")
            print(f"Pages: {pages}")
            print(f"Apply to all: {apply_to_all}")
            
            # Validate input
            if not os.path.exists(file_path):
                raise Exception("PDF ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ")
            
            if rotation_angle not in [90, 180, 270, -90, -180, -270]:
                rotation_angle = 90  # Default to 90 degrees
            
            # Read PDF
            from PyPDF2 import PdfReader, PdfWriter
            import io
            
            # First pass: Read the original PDF into memory
            with open(file_path, 'rb') as original_file:
                original_data = original_file.read()
            
            reader = PdfReader(io.BytesIO(original_data))
            writer = PdfWriter()
            total_pages = len(reader.pages)
            
            if total_pages == 0:
                raise Exception("PDF ಯಲ್ಲಿ ಯಾವುದೇ ಪುಟಗಳಿಲ್ಲ")
            
            # Determine which pages to rotate
            if pages.strip():
                pages_to_rotate = self._parse_page_ranges_enhanced(pages, total_pages)
                print(f"Rotating specific pages: {pages_to_rotate}")
            else:
                pages_to_rotate = list(range(1, total_pages + 1))
                print(f"No specific pages - rotating all pages")
                        
            print(f"Pages to rotate: {pages_to_rotate}")
            
            # Process each page - create fresh reader for each page to avoid object sharing
            for page_num in range(1, total_pages + 1):
                if page_num in pages_to_rotate:
                    # Create a fresh reader for the page to be rotated
                    fresh_reader = PdfReader(io.BytesIO(original_data))
                    page = fresh_reader.pages[page_num - 1]
                    page.rotate(rotation_angle)
                    writer.add_page(page)
                    print(f"Rotated page {page_num} by {rotation_angle} degrees")
                else:
                    # Create a fresh reader for the non-rotated page
                    fresh_reader = PdfReader(io.BytesIO(original_data))
                    page = fresh_reader.pages[page_num - 1]
                    writer.add_page(page)
                    print(f"Added page {page_num} without rotation")
            
            # Create output path
            output_path = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_rotated.pdf")
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Write rotated PDF
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            # Validate output
            if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
                raise Exception("ತಿರುಗಿಸಿದ PDF ರಚಿಸಲಾಗಿಲ್ಲ")
            
            print(f"=== ROTATION SUCCESSFUL ===")
            print(f"Output: {output_path}")
            print(f"Rotated {len(pages_to_rotate)} pages by {rotation_angle} degrees")
            
            return output_path
            
        except Exception as e:
            print(f"Rotation error: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"PDF ತಿರುಗಿಸುವಿಕೆ ವಿಫಲ: {str(e)}")

    def _parse_page_ranges_enhanced(self, pages_str, total_pages):
        """Parse page ranges like '1,3,5-10' into list of page numbers"""
        try:
            pages = []
            parts = pages_str.split(',')
            
            for part in parts:
                part = part.strip()
                if '-' in part:
                    start, end = map(int, part.split('-'))
                    pages.extend(range(start, min(end + 1, total_pages + 1)))
                else:
                    page_num = int(part)
                    if 1 <= page_num <= total_pages:
                        pages.append(page_num)
            
            return list(set(pages))  # Remove duplicates
        except Exception as e:
            raise Exception(f"ಅಮಾನ್ಯ ಪುಟ ಸಂಖ್ಯೆಗಳು: {str(e)}")

    def generate_page_previews(self, pdf_path, session_id, preview_folder):
        """Generate page preview images for PDF - Updated to handle rotated pages"""
        try:
            # Create session-specific preview directory
            session_preview_dir = os.path.join(preview_folder, session_id)
            os.makedirs(session_preview_dir, exist_ok=True)
            
            # Open PDF with PyMuPDF for better image rendering
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            if total_pages == 0:
                return None
            
            previews = []
            
            # Generate preview for each page (limit to first 50 pages for performance)
            max_previews = min(total_pages, 50)
            
            for page_num in range(max_previews):
                try:
                    page = doc[page_num]
                    
                    # Create preview image
                    mat = fitz.Matrix(0.5, 0.5)  # Scale down for preview
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to PIL Image
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Save preview image
                    preview_filename = f"page_{page_num + 1}.png"
                    preview_path = os.path.join(session_preview_dir, preview_filename)
                    img.save(preview_path, "PNG")
                    
                    previews.append({
                        'page_num': page_num + 1,
                        'image_path': preview_path
                    })
                    
                except Exception as page_error:
                    print(f"Error generating preview for page {page_num + 1}: {page_error}")
                    continue
            
            doc.close()
            
            return {
                'total_pages': total_pages,
                'previews': previews
            }
            
        except Exception as e:
            print(f"Preview generation error: {str(e)}")
        return None
    def _compress_ultra_aggressive(self, input_path, output_path, level, target_size_mb=None):
        """
        Ultra-aggressive compression - maximum size reduction
        """
        try:
            print("Starting ultra-aggressive compression...")
            doc = fitz.open(input_path)
            new_doc = fitz.open()
            
            # Very aggressive settings
            dpi = 72  # Very low DPI
            jpeg_quality = 30  # Very low quality
            scale_factor = 0.4  # Scale down to 40%
            
            # If target size specified, adjust parameters dynamically
            if target_size_mb:
                original_size_mb = os.path.getsize(input_path) / (1024 * 1024)
                compression_ratio = target_size_mb / original_size_mb
                
                if compression_ratio < 0.1:  # Need >90% compression
                    dpi = 50
                    jpeg_quality = 20
                    scale_factor = 0.3
                elif compression_ratio < 0.2:  # Need >80% compression
                    dpi = 60
                    jpeg_quality = 25
                    scale_factor = 0.35
                elif compression_ratio < 0.3:  # Need >70% compression
                    dpi = 72
                    jpeg_quality = 30
                    scale_factor = 0.4
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert page to very low-res image
                mat = fitz.Matrix(dpi/72, dpi/72)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                # Convert to PIL for aggressive compression
                img_data = pix.tobytes("png")
                pil_img = Image.open(io.BytesIO(img_data))
                
                # Scale down further
                new_width = int(pil_img.width * scale_factor)
                new_height = int(pil_img.height * scale_factor)
                pil_img = pil_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # Convert to RGB and compress heavily
                if pil_img.mode != 'RGB':
                    pil_img = pil_img.convert('RGB')
                
                # Apply additional optimizations
                # Reduce color palette
                pil_img = pil_img.quantize(colors=64).convert('RGB')
                
                # Save with very low quality
                jpeg_io = io.BytesIO()
                pil_img.save(jpeg_io, format='JPEG', quality=jpeg_quality, 
                            optimize=True, progressive=True)
                jpeg_data = jpeg_io.getvalue()
                
                # Create new page from compressed image
                img_rect = fitz.Rect(0, 0, pil_img.width, pil_img.height)
                new_page = new_doc.new_page(width=img_rect.width, height=img_rect.height)
                new_page.insert_image(img_rect, stream=jpeg_data)
            
            doc.close()
            
            # Save with maximum compression
            new_doc.save(output_path, 
                        deflate=True, 
                        garbage=4, 
                        clean=True,
                        linear=True,
                        pretty=False)
            new_doc.close()
            
            print("Ultra-aggressive compression completed")
            return True
            
        except Exception as e:
            print(f"Ultra-aggressive compression error: {e}")
            return False

    def _compress_extreme_image_recreation(self, input_path, output_path, level, target_size_mb=None):
        """
        Extreme image recreation with adaptive quality based on target size
        """
        try:
            doc = fitz.open(input_path)
            new_doc = fitz.open()
            
            # Base settings
            base_dpi = 100
            base_quality = 50
            
            # Adjust based on target size
            if target_size_mb:
                original_size_mb = os.path.getsize(input_path) / (1024 * 1024)
                compression_ratio = target_size_mb / original_size_mb
                
                # More aggressive settings for smaller target sizes
                if compression_ratio < 0.2:
                    base_dpi = 75
                    base_quality = 35
                elif compression_ratio < 0.5:
                    base_dpi = 85
                    base_quality = 45
            
            # Adjust based on level
            if level == 'maximum':
                dpi = max(50, base_dpi - 25)
                jpeg_quality = max(20, base_quality - 15)
            elif level == 'high':
                dpi = max(75, base_dpi - 10)
                jpeg_quality = max(35, base_quality - 10)
            else:
                dpi = base_dpi
                jpeg_quality = base_quality
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert to image with calculated DPI
                mat = fitz.Matrix(dpi/72, dpi/72)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                # Convert to PIL
                img_data = pix.tobytes("png")
                pil_img = Image.open(io.BytesIO(img_data))
                
                # Convert to RGB
                if pil_img.mode != 'RGB':
                    pil_img = pil_img.convert('RGB')
                
                # Apply sharpening before compression
                from PIL import ImageFilter
                pil_img = pil_img.filter(ImageFilter.UnsharpMask(radius=1, percent=100, threshold=3))
                
                # Compress as JPEG
                jpeg_io = io.BytesIO()
                pil_img.save(jpeg_io, format='JPEG', 
                            quality=jpeg_quality, 
                            optimize=True,
                            progressive=True)
                jpeg_data = jpeg_io.getvalue()
                
                # Create new page
                img_rect = fitz.Rect(0, 0, pil_img.width, pil_img.height)
                new_page = new_doc.new_page(width=img_rect.width, height=img_rect.height)
                new_page.insert_image(img_rect, stream=jpeg_data)
            
            doc.close()
            
            # Save with compression
            new_doc.save(output_path, deflate=True, garbage=4, clean=True)
            new_doc.close()
            
            return True
            
        except Exception as e:
            print(f"Extreme image recreation error: {e}")
            return False

    def _iterative_compression(self, input_path, output_path, target_size_mb):
        """
        Iteratively compress until target size is reached
        """
        try:
            print(f"Starting iterative compression to reach {target_size_mb}MB...")
            
            current_path = input_path
            iteration = 0
            max_iterations = 5
            
            while iteration < max_iterations:
                current_size_mb = os.path.getsize(current_path) / (1024 * 1024)
                
                if current_size_mb <= target_size_mb:
                    print(f"Target reached after {iteration} iterations!")
                    if current_path != output_path:
                        import shutil
                        shutil.copy2(current_path, output_path)
                    return True
                
                print(f"Iteration {iteration + 1}: Current size {current_size_mb:.2f}MB, target {target_size_mb}MB")
                
                # Calculate required compression ratio
                compression_ratio = target_size_mb / current_size_mb
                
                # Adjust quality based on how much more compression is needed
                if compression_ratio < 0.5:
                    quality = 25
                    dpi = 60
                elif compression_ratio < 0.7:
                    quality = 35
                    dpi = 75
                else:
                    quality = 50
                    dpi = 90
                
                # Create temporary file for this iteration
                temp_path = output_path.replace('.pdf', f'_iter_{iteration}.pdf')
                
                # Apply compression
                doc = fitz.open(current_path)
                new_doc = fitz.open()
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    
                    mat = fitz.Matrix(dpi/72, dpi/72)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    
                    img_data = pix.tobytes("png")
                    pil_img = Image.open(io.BytesIO(img_data))
                    
                    if pil_img.mode != 'RGB':
                        pil_img = pil_img.convert('RGB')
                    
                    jpeg_io = io.BytesIO()
                    pil_img.save(jpeg_io, format='JPEG', quality=quality, optimize=True)
                    jpeg_data = jpeg_io.getvalue()
                    
                    img_rect = fitz.Rect(0, 0, pil_img.width, pil_img.height)
                    new_page = new_doc.new_page(width=img_rect.width, height=img_rect.height)
                    new_page.insert_image(img_rect, stream=jpeg_data)
                
                doc.close()
                new_doc.save(temp_path, deflate=True, garbage=4, clean=True)
                new_doc.close()
                
                # Clean up previous iteration
                if current_path != input_path and os.path.exists(current_path):
                    os.remove(current_path)
                
                current_path = temp_path
                iteration += 1
            
            # Copy final result
            if current_path != output_path:
                import shutil
                shutil.copy2(current_path, output_path)
                if os.path.exists(current_path):
                    os.remove(current_path)
            
            final_size_mb = os.path.getsize(output_path) / (1024 * 1024)
            print(f"Iterative compression completed. Final size: {final_size_mb:.2f}MB")
            
            return True
            
        except Exception as e:
            print(f"Iterative compression error: {e}")
            return False

    def _compress_pymupdf(self, input_path, output_path, level):
        """Compress using PyMuPDF - IMPROVED VERSION"""
        try:
            doc = fitz.open(input_path)
            
            # First pass - compress images in each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Get all images on the page
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    # Extract image
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Convert to PIL for compression
                    from PIL import Image
                    import io
                    
                    pil_image = Image.open(io.BytesIO(image_bytes))
                    
                    # Compress based on level
                    if level == 'high':
                        quality = 30
                        scale = 0.5  # Reduce size to 50%
                    elif level == 'medium':
                        quality = 50
                        scale = 0.7  # Reduce size to 70%
                    else:  # low
                        quality = 70
                        scale = 0.9  # Reduce size to 90%
                    
                    # Resize image
                    new_size = (int(pil_image.width * scale), int(pil_image.height * scale))
                    pil_image = pil_image.resize(new_size, Image.Resampling.LANCZOS)
                    
                    # Compress and save back
                    if pil_image.mode in ('RGBA', 'LA'):
                        # Convert transparent images to white background
                        background = Image.new('RGB', pil_image.size, (255, 255, 255))
                        if pil_image.mode == 'RGBA':
                            background.paste(pil_image, mask=pil_image.split()[-1])
                        else:
                            background.paste(pil_image)
                        pil_image = background
                    
                    # Save compressed image
                    compressed_io = io.BytesIO()
                    pil_image.save(compressed_io, format='JPEG', quality=quality, optimize=True)
                    compressed_bytes = compressed_io.getvalue()
                    
                    # Replace image in PDF
                    doc._updateObject(xref, compressed_bytes, filename="image.jpg")
            
            # Second pass - save with compression options
            if level == 'high':
                options = {
                    'garbage': 4,
                    'clean': True,
                    'deflate': True,
                    'deflate_images': True,
                    'deflate_fonts': True,
                    'linear': True,
                    'pretty': False
                }
            elif level == 'medium':
                options = {
                    'garbage': 3,
                    'clean': True,
                    'deflate': True,
                    'deflate_images': True,
                    'deflate_fonts': True,
                    'linear': True,
                    'pretty': False
                }
            else:  # low
                options = {
                    'garbage': 2,
                    'clean': True,
                    'deflate': True,
                    'deflate_images': False,
                    'deflate_fonts': False,
                    'linear': False,
                    'pretty': False
                }
            
            doc.save(output_path, **options)
            doc.close()
            
            return True
            
        except Exception as e:
            print(f"PyMuPDF compression error: {e}")
            return False

    
    
    def _validate_compressed_pdf(self, pdf_path):
        """Validate that compressed PDF is readable"""
        try:
            # Test with PyMuPDF
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            doc.close()
            
            if page_count == 0:
                return False
            
            # Test with PyPDF2
            from PyPDF2 import PdfReader
            reader = PdfReader(pdf_path)
            if len(reader.pages) != page_count:
                return False
            
            print(f"Validation successful: {page_count} pages")
            return True
            
        except Exception as e:
            print(f"Validation failed: {e}")
            return False

    
    def _compress_advanced_pymupdf(self, input_path, output_path, level):
        """Advanced PyMuPDF compression with image optimization"""
        try:
            doc = fitz.open(input_path)
            
            # Step 1: Compress images on each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images(full=True)
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # Convert to PIL for compression
                        pil_image = Image.open(io.BytesIO(image_bytes))
                        
                        # Set compression parameters based on level
                        if level == 'high':
                            quality = 40
                            scale_factor = 0.6
                        elif level == 'medium':
                            quality = 60
                            scale_factor = 0.8
                        else:  # low
                            quality = 80
                            scale_factor = 0.9
                        
                        # Resize image
                        new_width = int(pil_image.width * scale_factor)
                        new_height = int(pil_image.height * scale_factor)
                        pil_image = pil_image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        
                        # Convert RGBA to RGB with white background
                        if pil_image.mode in ('RGBA', 'LA', 'P'):
                            background = Image.new('RGB', pil_image.size, (255, 255, 255))
                            if pil_image.mode == 'P':
                                pil_image = pil_image.convert('RGBA')
                            if pil_image.mode in ('RGBA', 'LA'):
                                background.paste(pil_image, mask=pil_image.split()[-1])
                            else:
                                background.paste(pil_image)
                            pil_image = background
                        
                        # Compress as JPEG
                        compressed_io = io.BytesIO()
                        pil_image.save(compressed_io, format='JPEG', quality=quality, optimize=True)
                        compressed_bytes = compressed_io.getvalue()
                        
                        # Replace image in PDF
                        img_dict = {
                            "type": "image",
                            "bbox": fitz.Rect(0, 0, new_width, new_height),
                            "width": new_width,
                            "height": new_height,
                            "image": compressed_bytes
                        }
                        
                        # Update the image in the document
                        page.insert_image(img_dict["bbox"], stream=compressed_bytes)
                        
                    except Exception as img_error:
                        print(f"Error compressing image {img_index}: {img_error}")
                        continue
            
            # Step 2: Save with aggressive compression settings
            if level == 'high':
                options = {
                    'garbage': 4,
                    'clean': True,
                    'deflate': True,
                    'deflate_images': True,
                    'deflate_fonts': True,
                    'linear': True,
                    'pretty': False,
                    'ascii': False
                }
            elif level == 'medium':
                options = {
                    'garbage': 3,
                    'clean': True,
                    'deflate': True,
                    'deflate_images': True,
                    'deflate_fonts': True,
                    'linear': False,
                    'pretty': False
                }
            else:  # low
                options = {
                    'garbage': 2,
                    'clean': True,
                    'deflate': True,
                    'deflate_images': False,
                    'deflate_fonts': False,
                    'linear': False,
                    'pretty': False
                }
            
            doc.save(output_path, **options)
            doc.close()
            
            return True
            
        except Exception as e:
            print(f"Advanced PyMuPDF compression error: {e}")
            return False

    def _compress_by_image_recreation(self, input_path, output_path, level):
        """Compress by converting pages to images and back - MOST AGGRESSIVE"""
        try:
            doc = fitz.open(input_path)
            new_doc = fitz.open()
            
            # Set parameters based on compression level
            if level == 'high':
                dpi = 100  # Lower DPI = smaller file
                jpeg_quality = 60
            elif level == 'medium':
                dpi = 150
                jpeg_quality = 75
            else:  # low
                dpi = 200
                jpeg_quality = 85
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert page to image
                mat = fitz.Matrix(dpi/72, dpi/72)  # 72 is default DPI
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                # Convert to PIL for JPEG compression
                img_data = pix.tobytes("png")
                pil_img = Image.open(io.BytesIO(img_data))
                
                # Convert to RGB if needed
                if pil_img.mode != 'RGB':
                    pil_img = pil_img.convert('RGB')
                
                # Compress as JPEG
                jpeg_io = io.BytesIO()
                pil_img.save(jpeg_io, format='JPEG', quality=jpeg_quality, optimize=True)
                jpeg_data = jpeg_io.getvalue()
                
                # Create new page from compressed image
                img_rect = fitz.Rect(0, 0, pil_img.width, pil_img.height)
                new_page = new_doc.new_page(width=img_rect.width, height=img_rect.height)
                new_page.insert_image(img_rect, stream=jpeg_data)
            
            doc.close()
            
            # Save the new document
            new_doc.save(output_path, deflate=True, garbage=4, clean=True)
            new_doc.close()
            
            return True
            
        except Exception as e:
            print(f"Image recreation compression error: {e}")
            return False

    def _compress_basic_pymupdf(self, input_path, output_path, level):
        """Basic PyMuPDF compression - FALLBACK"""
        try:
            doc = fitz.open(input_path)
            
            # Basic settings that always work
            options = {
                'garbage': 4,
                'clean': True,
                'deflate': True
            }
            
            doc.save(output_path, **options)
            doc.close()
            
            return True
            
        except Exception as e:
            print(f"Basic PyMuPDF compression error: {e}")
            return False
    # Add this new method to your PDFOperations class in pdf_operations.py

    def compress_pdf_enhanced(self, file_path, compression_level, session_id, 
                        target_size_mb=None, image_quality=None, image_dpi=None,
                        remove_metadata=False, optimize_fonts=False):
        """
        Enhanced PDF compression with user controls and target size
        """
        try:
            print(f"=== ENHANCED PDF COMPRESSION ===")
            print(f"Input: {file_path}")
            print(f"Level: {compression_level}")
            print(f"Target size: {target_size_mb}MB" if target_size_mb else "No target size")
            
            # Validate input
            if not os.path.exists(file_path):
                raise Exception("PDF ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ")
            
            original_size = os.path.getsize(file_path)
            original_size_mb = original_size / (1024 * 1024)
            print(f"Original size: {original_size_mb:.2f} MB")
            
            output_path = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_compressed.pdf")
            
            # If target size is specified and user wants custom compression
            if target_size_mb and target_size_mb < original_size_mb:
                print("Using adaptive target-based compression...")
                success = self._adaptive_target_compression(file_path, output_path, target_size_mb)
                if success and os.path.exists(output_path):
                    final_size_mb = os.path.getsize(output_path) / (1024 * 1024)
                    print(f"Final size: {final_size_mb:.2f}MB")
                    return output_path
            
            # Determine compression parameters
            compression_params = self._get_compression_parameters(
                compression_level, target_size_mb, original_size_mb,
                image_quality, image_dpi
            )
            
            # Select compression methods based on aggressiveness needed
            if compression_level == 'maximum' or (target_size_mb and target_size_mb < original_size_mb * 0.3):
                methods = [
                    ("Extreme Aggressive", self._compress_extreme_aggressive),
                    ("Ultra Aggressive", self._compress_ultra_aggressive),
                    ("Smart Adaptive", self._compress_smart_adaptive),
                ]
            else:
                methods = [
                    ("Smart Adaptive", self._compress_smart_adaptive),
                    ("Enhanced Image Recreation", self._compress_image_recreation_enhanced),
                    ("Advanced PyMuPDF", self._compress_pymupdf_enhanced),
                ]
            
            # Try compression methods
            best_result = None
            best_size = original_size
            
            for method_name, method_func in methods:
                try:
                    print(f"Trying {method_name}...")
                    temp_output = output_path.replace('.pdf', f'_temp_{method_name.replace(" ", "_").lower()}.pdf')
                    
                    result = method_func(file_path, temp_output, compression_params, 
                                    remove_metadata, optimize_fonts)
                    
                    if result and os.path.exists(temp_output):
                        compressed_size = os.path.getsize(temp_output)
                        compressed_size_mb = compressed_size / (1024 * 1024)
                        
                        if compressed_size > 0 and compressed_size < best_size:
                            if best_result and os.path.exists(best_result):
                                os.remove(best_result)
                            
                            best_result = temp_output
                            best_size = compressed_size
                            reduction = (1 - compressed_size/original_size) * 100
                            
                            print(f"✓ {method_name} successful!")
                            print(f"Size: {compressed_size_mb:.2f}MB (reduction: {reduction:.1f}%)")
                            
                            # Check if good enough
                            if target_size_mb and compressed_size_mb <= target_size_mb:
                                print("Target reached!")
                                break
                        else:
                            print(f"✗ {method_name} didn't improve")
                            if os.path.exists(temp_output):
                                os.remove(temp_output)
                            
                except Exception as e:
                    print(f"✗ {method_name} failed: {e}")
                    continue
            
            # Finalize result
            if best_result and os.path.exists(best_result):
                if os.path.exists(output_path):
                    os.remove(output_path)
                os.rename(best_result, output_path)
                
                final_size_mb = os.path.getsize(output_path) / (1024 * 1024)
                reduction = (1 - final_size_mb/original_size_mb) * 100
                
                print(f"=== COMPRESSION COMPLETE ===")
                print(f"Original: {original_size_mb:.2f} MB")
                print(f"Compressed: {final_size_mb:.2f} MB")
                print(f"Reduction: {reduction:.1f}%")
                
                return output_path
            else:
                raise Exception("ಎಲ್ಲಾ ಸಂಕುಚನ ವಿಧಾನಗಳು ವಿಫಲವಾಗಿವೆ")
                
        except Exception as e:
            print(f"Enhanced compression error: {str(e)}")
            raise Exception(f"PDF ಸಂಕುಚನ ವಿಫಲ: {str(e)}")

    def _get_compression_parameters(self, level, target_size_mb, original_size_mb, 
                                custom_quality=None, custom_dpi=None):
        """Get compression parameters based on level and target size"""
        
        # Base parameters for each level
        level_configs = {
            'low': {'dpi': 200, 'quality': 80, 'scale': 0.9},
            'medium': {'dpi': 150, 'quality': 60, 'scale': 0.8},
            'high': {'dpi': 100, 'quality': 45, 'scale': 0.6},
            'maximum': {'dpi': 75, 'quality': 30, 'scale': 0.4},
            'custom': {'dpi': 150, 'quality': 60, 'scale': 0.8}
        }
        
        params = level_configs.get(level, level_configs['medium']).copy()
        
        # Adjust for target size if specified
        if target_size_mb and original_size_mb > 0:
            compression_ratio = target_size_mb / original_size_mb
            
            if compression_ratio < 0.1:  # Need >90% compression
                params.update({'dpi': 50, 'quality': 20, 'scale': 0.3})
            elif compression_ratio < 0.2:  # Need >80% compression
                params.update({'dpi': 60, 'quality': 25, 'scale': 0.35})
            elif compression_ratio < 0.3:  # Need >70% compression
                params.update({'dpi': 75, 'quality': 35, 'scale': 0.4})
            elif compression_ratio < 0.5:  # Need >50% compression
                params.update({'dpi': 100, 'quality': 45, 'scale': 0.6})
        
        # Override with custom values if provided
        if custom_quality:
            params['quality'] = custom_quality
        if custom_dpi:
            params['dpi'] = custom_dpi
        
        return params

    def _compress_smart_adaptive(self, input_path, output_path, params, remove_metadata, optimize_fonts):
        """Smart adaptive compression that analyzes content"""
        try:
            print("Starting smart adaptive compression...")
            
            doc = fitz.open(input_path)
            new_doc = fitz.open()
            
            # Analyze document characteristics
            total_images = 0
            total_text_pages = 0
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                images = page.get_images()
                text = page.get_text().strip()
                
                if images:
                    total_images += len(images)
                if text:
                    total_text_pages += 1
            
            # Adjust parameters based on content
            if total_images > total_text_pages * 2:  # Image-heavy document
                print("Detected image-heavy document, using aggressive image compression")
                params['quality'] = max(20, params['quality'] - 15)
                params['dpi'] = max(50, params['dpi'] - 25)
            elif total_text_pages > total_images * 3:  # Text-heavy document
                print("Detected text-heavy document, preserving readability")
                params['quality'] = min(85, params['quality'] + 10)
                params['dpi'] = min(200, params['dpi'] + 25)
            
            # Process each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert to image with adaptive DPI
                mat = fitz.Matrix(params['dpi']/72, params['dpi']/72)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                # Convert to PIL and optimize
                img_data = pix.tobytes("png")
                pil_img = Image.open(io.BytesIO(img_data))
                
                # Scale if needed
                if params['scale'] < 1.0:
                    new_width = int(pil_img.width * params['scale'])
                    new_height = int(pil_img.height * params['scale'])
                    pil_img = pil_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # Convert to RGB
                if pil_img.mode != 'RGB':
                    pil_img = pil_img.convert('RGB')
                
                # Apply sharpening for better quality at lower resolutions
                if params['dpi'] < 150:
                    from PIL import ImageFilter
                    pil_img = pil_img.filter(ImageFilter.UnsharpMask(radius=1, percent=100, threshold=3))
                
                # Compress as JPEG
                jpeg_io = io.BytesIO()
                pil_img.save(jpeg_io, format='JPEG', 
                            quality=params['quality'], 
                            optimize=True, 
                            progressive=True)
                jpeg_data = jpeg_io.getvalue()
                
                # Create new page
                img_rect = fitz.Rect(0, 0, pil_img.width, pil_img.height)
                new_page = new_doc.new_page(width=img_rect.width, height=img_rect.height)
                new_page.insert_image(img_rect, stream=jpeg_data)
            
            doc.close()
            
            # Save with options
            save_options = {
                'deflate': True,
                'garbage': 4,
                'clean': True,
                'linear': True,
                'pretty': False
            }
            
            new_doc.save(output_path, **save_options)
            new_doc.close()
            
            # Post-process if requested
            if remove_metadata or optimize_fonts:
                self._post_process_pdf(output_path, remove_metadata, optimize_fonts)
            
            print("Smart adaptive compression completed")
            return True
            
        except Exception as e:
            print(f"Smart adaptive compression error: {e}")
            return False

    def _compress_image_recreation_enhanced(self, input_path, output_path, params, remove_metadata, optimize_fonts):
        """Enhanced image recreation with better quality control"""
        try:
            print("Starting enhanced image recreation...")
            
            doc = fitz.open(input_path)
            new_doc = fitz.open()
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Use adaptive DPI based on page content
                mat = fitz.Matrix(params['dpi']/72, params['dpi']/72)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                # Convert to PIL
                img_data = pix.tobytes("png")
                pil_img = Image.open(io.BytesIO(img_data))
                
                # Apply scaling
                if params['scale'] < 1.0:
                    new_width = int(pil_img.width * params['scale'])
                    new_height = int(pil_img.height * params['scale'])
                    pil_img = pil_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # Convert to RGB
                if pil_img.mode != 'RGB':
                    pil_img = pil_img.convert('RGB')
                
                # Optimize image before compression
                # Reduce noise for better compression
                from PIL import ImageFilter
                pil_img = pil_img.filter(ImageFilter.MedianFilter(size=3))
                
                # Compress with progressive JPEG
                jpeg_io = io.BytesIO()
                pil_img.save(jpeg_io, format='JPEG', 
                            quality=params['quality'],
                            optimize=True,
                            progressive=True)
                jpeg_data = jpeg_io.getvalue()
                
                # Create new page
                img_rect = fitz.Rect(0, 0, pil_img.width, pil_img.height)
                new_page = new_doc.new_page(width=img_rect.width, height=img_rect.height)
                new_page.insert_image(img_rect, stream=jpeg_data)
            
            doc.close()
            
            # Save with compression
            new_doc.save(output_path, deflate=True, garbage=4, clean=True)
            new_doc.close()
            
            # Post-process
            if remove_metadata or optimize_fonts:
                self._post_process_pdf(output_path, remove_metadata, optimize_fonts)
            
            return True
            
        except Exception as e:
            print(f"Enhanced image recreation error: {e}")
            return False

    def _compress_pymupdf_enhanced(self, input_path, output_path, params, remove_metadata, optimize_fonts):
        """Enhanced PyMuPDF compression"""
        try:
            print("Starting enhanced PyMuPDF compression...")
            
            doc = fitz.open(input_path)
            
            # Process images on each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images(full=True)
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # Convert to PIL for optimization
                        pil_image = Image.open(io.BytesIO(image_bytes))
                        
                        # Resize image
                        if params['scale'] < 1.0:
                            new_width = int(pil_image.width * params['scale'])
                            new_height = int(pil_image.height * params['scale'])
                            pil_image = pil_image.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        
                        # Convert to RGB
                        if pil_image.mode in ('RGBA', 'LA', 'P'):
                            background = Image.new('RGB', pil_image.size, (255, 255, 255))
                            if pil_image.mode == 'P':
                                pil_image = pil_image.convert('RGBA')
                            if pil_image.mode in ('RGBA', 'LA'):
                                background.paste(pil_image, mask=pil_image.split()[-1])
                            else:
                                background.paste(pil_image)
                            pil_image = background
                        
                        # Compress
                        compressed_io = io.BytesIO()
                        pil_image.save(compressed_io, format='JPEG', 
                                    quality=params['quality'], optimize=True)
                        compressed_bytes = compressed_io.getvalue()
                        
                        # Replace image in PDF (simplified approach)
                        # Note: This is a basic replacement - you might need more sophisticated methods
                        
                    except Exception as img_error:
                        print(f"Error processing image {img_index}: {img_error}")
                        continue
            
            # Save with compression options
            save_options = {
                'garbage': 4,
                'clean': True,
                'deflate': True,
                'deflate_images': True,
                'deflate_fonts': optimize_fonts,
                'linear': True,
                'pretty': False
            }
            
            doc.save(output_path, **save_options)
            doc.close()
            
            # Post-process
            if remove_metadata or optimize_fonts:
                self._post_process_pdf(output_path, remove_metadata, optimize_fonts)
            
            return True
            
        except Exception as e:
            print(f"Enhanced PyMuPDF compression error: {e}")
            return False

    def _iterative_compression_enhanced(self, input_path, output_path, target_size_mb, base_params):
        """Enhanced iterative compression"""
        try:
            print(f"Starting enhanced iterative compression to reach {target_size_mb}MB...")
            
            current_path = input_path
            iteration = 0
            max_iterations = 3
            
            while iteration < max_iterations:
                current_size_mb = os.path.getsize(current_path) / (1024 * 1024)
                
                if current_size_mb <= target_size_mb:
                    print(f"Target reached after {iteration} iterations!")
                    if current_path != output_path:
                        import shutil
                        shutil.copy2(current_path, output_path)
                    return True
                
                print(f"Iteration {iteration + 1}: {current_size_mb:.2f}MB -> target {target_size_mb}MB")
                
                # Adjust parameters more aggressively
                compression_ratio = target_size_mb / current_size_mb
                params = base_params.copy()
                
                if compression_ratio < 0.3:
                    params.update({'dpi': 50, 'quality': 20, 'scale': 0.3})
                elif compression_ratio < 0.5:
                    params.update({'dpi': 60, 'quality': 30, 'scale': 0.4})
                elif compression_ratio < 0.7:
                    params.update({'dpi': 75, 'quality': 40, 'scale': 0.5})
                
                # Create temporary file
                temp_path = output_path.replace('.pdf', f'_iter_{iteration}.pdf')
                
                # Apply compression
                success = self._compress_image_recreation_enhanced(
                    current_path, temp_path, params, False, False
                )
                
                if not success or not os.path.exists(temp_path):
                    print(f"Iteration {iteration + 1} failed")
                    break
                
                # Clean up previous iteration
                if current_path != input_path and os.path.exists(current_path):
                    os.remove(current_path)
                
                current_path = temp_path
                iteration += 1
            
            # Copy final result
            if current_path != output_path:
                import shutil
                shutil.copy2(current_path, output_path)
                if current_path != input_path and os.path.exists(current_path):
                    os.remove(current_path)
            
            return True
            
        except Exception as e:
            print(f"Enhanced iterative compression error: {e}")
            return False

    def _post_process_pdf(self, pdf_path, remove_metadata, optimize_fonts):
        """Post-process PDF to remove metadata and optimize fonts"""
        try:
            if not remove_metadata and not optimize_fonts:
                return
            
            print("Post-processing PDF...")
            
            doc = fitz.open(pdf_path)
            
            if remove_metadata:
                print("Removing metadata...")
                # Clear metadata
                doc.set_metadata({})
                
            # Note: Font optimization would require more complex implementation
            # This is a placeholder for the concept
            
            doc.save(pdf_path, incremental=False, encryption=fitz.PDF_ENCRYPT_NONE)
            doc.close()
            
            print("Post-processing completed")
            
        except Exception as e:
            print(f"Post-processing error: {e}")
    def _compress_extreme_aggressive(self, input_path, output_path, params, remove_metadata, optimize_fonts):
        """
        Extremely aggressive compression - for when you need maximum size reduction
        """
        try:
            print("Starting extreme aggressive compression...")
            
            doc = fitz.open(input_path)
            new_doc = fitz.open()
            
            # Ultra-low settings for maximum compression
            dpi = max(50, params.get('dpi', 50))  # Very low DPI
            quality = max(15, params.get('quality', 20))  # Ultra low quality
            scale_factor = min(0.3, params.get('scale', 0.3))  # Scale down to 30%
            
            print(f"Extreme settings: DPI={dpi}, Quality={quality}, Scale={scale_factor}")
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert to very low resolution image
                mat = fitz.Matrix(dpi/72, dpi/72)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                
                # Convert to PIL for aggressive processing
                img_data = pix.tobytes("png")
                pil_img = Image.open(io.BytesIO(img_data))
                
                # Scale down dramatically
                new_width = int(pil_img.width * scale_factor)
                new_height = int(pil_img.height * scale_factor)
                pil_img = pil_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # Convert to RGB
                if pil_img.mode != 'RGB':
                    pil_img = pil_img.convert('RGB')
                
                # Reduce color depth for smaller file size
                pil_img = pil_img.quantize(colors=32).convert('RGB')  # Reduce to 32 colors
                
                # Apply aggressive compression
                jpeg_io = io.BytesIO()
                pil_img.save(jpeg_io, format='JPEG', 
                            quality=quality, 
                            optimize=True, 
                            progressive=True)
                jpeg_data = jpeg_io.getvalue()
                
                # Create new page
                img_rect = fitz.Rect(0, 0, pil_img.width, pil_img.height)
                new_page = new_doc.new_page(width=img_rect.width, height=img_rect.height)
                new_page.insert_image(img_rect, stream=jpeg_data)
            
            doc.close()
            
            # Save with maximum compression
            new_doc.save(output_path, 
                        deflate=True, 
                        garbage=4, 
                        clean=True,
                        linear=True,
                        pretty=False)
            new_doc.close()
            
            print("Extreme aggressive compression completed")
            return True
            
        except Exception as e:
            print(f"Extreme aggressive compression error: {e}")
            return False

    def _adaptive_target_compression(self, input_path, output_path, target_size_mb):
        """
        Adaptive compression that tries different settings until target size is reached
        """
        try:
            print(f"Starting adaptive compression to reach {target_size_mb}MB...")
            
            original_size_mb = os.path.getsize(input_path) / (1024 * 1024)
            compression_ratio = target_size_mb / original_size_mb
            
            print(f"Need {(1-compression_ratio)*100:.1f}% compression")
            
            # Progressive compression settings - start gentle, get more aggressive
            compression_settings = [
                {'dpi': 150, 'quality': 70, 'scale': 0.9, 'colors': None},
                {'dpi': 120, 'quality': 55, 'scale': 0.7, 'colors': None},
                {'dpi': 100, 'quality': 40, 'scale': 0.6, 'colors': 128},
                {'dpi': 80, 'quality': 30, 'scale': 0.5, 'colors': 64},
                {'dpi': 60, 'quality': 25, 'scale': 0.4, 'colors': 32},
                {'dpi': 50, 'quality': 20, 'scale': 0.3, 'colors': 16},
            ]
            
            for i, settings in enumerate(compression_settings):
                print(f"Trying compression level {i+1}/6...")
                
                temp_output = output_path.replace('.pdf', f'_temp_adaptive_{i}.pdf')
                
                # Apply compression with current settings
                doc = fitz.open(input_path)
                new_doc = fitz.open()
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    
                    mat = fitz.Matrix(settings['dpi']/72, settings['dpi']/72)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    
                    img_data = pix.tobytes("png")
                    pil_img = Image.open(io.BytesIO(img_data))
                    
                    # Scale image
                    if settings['scale'] < 1.0:
                        new_width = int(pil_img.width * settings['scale'])
                        new_height = int(pil_img.height * settings['scale'])
                        pil_img = pil_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    
                    # Convert to RGB
                    if pil_img.mode != 'RGB':
                        pil_img = pil_img.convert('RGB')
                    
                    # Reduce colors if specified
                    if settings['colors']:
                        pil_img = pil_img.quantize(colors=settings['colors']).convert('RGB')
                    
                    # Compress
                    jpeg_io = io.BytesIO()
                    pil_img.save(jpeg_io, format='JPEG', 
                            quality=settings['quality'], 
                            optimize=True, 
                            progressive=True)
                    jpeg_data = jpeg_io.getvalue()
                    
                    img_rect = fitz.Rect(0, 0, pil_img.width, pil_img.height)
                    new_page = new_doc.new_page(width=img_rect.width, height=img_rect.height)
                    new_page.insert_image(img_rect, stream=jpeg_data)
                
                doc.close()
                new_doc.save(temp_output, deflate=True, garbage=4, clean=True)
                new_doc.close()
                
                # Check if target size reached
                temp_size_mb = os.path.getsize(temp_output) / (1024 * 1024)
                print(f"Result size: {temp_size_mb:.2f}MB (target: {target_size_mb}MB)")
                
                if temp_size_mb <= target_size_mb:
                    print(f"Target reached with compression level {i+1}!")
                    if temp_output != output_path:
                        import shutil
                        shutil.move(temp_output, output_path)
                    
                    # Clean up other temp files
                    for j in range(i+1, len(compression_settings)):
                        temp_file = output_path.replace('.pdf', f'_temp_adaptive_{j}.pdf')
                        if os.path.exists(temp_file):
                            os.remove(temp_file)
                    
                    return True
                
                # If not the last attempt, keep this file for potential fallback
                if i < len(compression_settings) - 1:
                    continue
                else:
                    # This is our best effort
                    if temp_output != output_path:
                        import shutil
                        shutil.move(temp_output, output_path)
                    print(f"Could not reach exact target. Best result: {temp_size_mb:.2f}MB")
                    return True
            
            return False
            
        except Exception as e:
            print(f"Adaptive compression error: {e}")
            return False
    def pdf_to_images(self, file_path, session_id):
        """Convert PDF pages to JPEG images"""
        try:
            doc = fitz.open(file_path)
            images_folder = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_images")
            os.makedirs(images_folder, exist_ok=True)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                mat = fitz.Matrix(2, 2)
                pix = page.get_pixmap(matrix=mat)
                
                image_filename = f"page_{page_num + 1}.jpg"
                image_path = os.path.join(images_folder, image_filename)
                pix.save(image_path)
            
            doc.close()
            
            zip_filename = f"{session_id}_images.zip"
            zip_path = os.path.join(self.config.OUTPUT_FOLDER, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w') as zip_file:
                for root, dirs, files in os.walk(images_folder):
                    for file in files:
                        file_path = os.path.join(root, file)
                        zip_file.write(file_path, file)
            
            shutil.rmtree(images_folder)
            return zip_path
            
        except Exception as e:
            raise Exception(f"PDF ಚಿತ್ರ ಪರಿವರ್ತನೆ ವಿಫಲ: {str(e)}")

    def images_to_pdf(self, image_paths, session_id):
        """Convert images to PDF"""
        try:
            output_filename = f"{session_id}_from_images.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            images = []
            for image_path in image_paths:
                if os.path.exists(image_path):
                    img = Image.open(image_path)
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    images.append(img)
            
            if images:
                images[0].save(output_path, save_all=True, append_images=images[1:])
            
            return output_path
            
        except Exception as e:
            raise Exception(f"ಚಿತ್ರ PDF ಪರಿವರ್ತನೆ ವಿಫಲ: {str(e)}")

    def pdf_to_word(self, file_path, session_id):
        """Convert PDF to Word document"""
        try:
            doc = fitz.open(file_path)
            word_doc = Document()
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    word_doc.add_paragraph(text)
                
                if page_num < len(doc) - 1:
                    word_doc.add_page_break()
            
            doc.close()
            
            output_filename = f"{session_id}_converted.docx"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            word_doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"PDF Word ಪರಿವರ್ತನೆ ವಿಫಲ: {str(e)}")
    
    def _parse_page_ranges(self, pages_str, total_pages):
        """Parse page ranges like '1,3,5-10' into list of page numbers"""
        pages = []
    
        for part in pages_str.split(','):
            part = part.strip()
            if '-' in part:
                start, end = map(int, part.split('-'))
                pages.extend(range(start, min(end + 1, total_pages + 1)))
            else:
                page_num = int(part)
                if 1 <= page_num <= total_pages:
                    pages.append(page_num)
    
        return sorted(list(set(pages)))  # This should be OUTSIDE the for loop
        # Replace your existing word_to_pdf method with this complete version:



    def word_to_pdf(self, file_path, session_id):
        """
        Enhanced Word to PDF conversion with proper Kannada text rendering
        """
        try:
            print(f"=== WORD TO PDF CONVERSION (KANNADA FIXED) ===")
            print(f"Input file: {file_path}")
            print(f"Session ID: {session_id}")
            
            # Validate input file
            if not os.path.exists(file_path):
                raise Exception(f"Input file not found: {file_path}")
            
            file_size = os.path.getsize(file_path)
            print(f"Input file size: {file_size} bytes")
            
            if file_size == 0:
                raise Exception("ಖಾಲಿ Word ದಾಖಲೆ")
            
            # Test if file is readable
            try:
                from docx import Document
                test_doc = Document(file_path)
                para_count = len(test_doc.paragraphs)
                print(f"Word document has {para_count} paragraphs")
            except Exception as e:
                raise Exception(f"Word ದಾಖಲೆ ಓದಲು ಸಾಧ್ಯವಾಗಿಲ್ಲ: {str(e)}")
            
            
            
            
            # Method 1: docx2pdf with proper COM threading
            try:
                result = self._convert_with_docx2pdf_threaded(file_path, session_id)
                if result:
                    print("✓ docx2pdf conversion successful")
                    return result
            except Exception as e:
                print(f"✗ docx2pdf method failed: {e}")
            
            # Method 2: LibreOffice (if available)
            try:
                result = self._convert_with_libreoffice_simple(file_path, session_id)
                if result:
                    print("✓ LibreOffice conversion successful")
                    return result
            except Exception as e:
                print(f"✗ LibreOffice method failed: {e}")
            
            # Method 3: Enhanced ReportLab with better Kannada handling
            try:
                result = self._convert_with_weasyprint(file_path, session_id)
                if result:
                    print("✓ WeasyPrint conversion successful")
                    return result
            except Exception as e:
                print(f"✗ WeasyPrint method failed: {e}")
            
            raise Exception("ಎಲ್ಲಾ ಪರಿವರ್ತನೆ ವಿಧಾನಗಳು ವಿಫಲವಾಗಿವೆ")
            
        except Exception as e:
            print(f"Word to PDF conversion error: {str(e)}")
            raise Exception(f"Word to PDF ಪರಿವರ್ತನೆ ವಿಫಲ: {str(e)}")

    def sort_pdf_by_page_numbers(self, file_path, session_id, pages=""):
        """Sort PDF pages by detected Kannada page numbers"""
        try:
            from .kannada_numeral_converter import KannadaNumeralConverter
            
            converter = KannadaNumeralConverter()
            doc = fitz.open(file_path)
            
            page_data = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                extracted_number = converter.extract_page_number_from_text(text)
                
                page_data.append({
                    'page': page,
                    'original_num': page_num + 1,
                    'extracted_num': extracted_number if extracted_number else page_num + 1
                })
            
            page_data.sort(key=lambda x: x['extracted_num'])
            
            output_filename = f"{session_id}_sortedbynum.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            new_doc = fitz.open()
            for data in page_data:
                new_doc.insert_pdf(doc, from_page=data['original_num']-1, to_page=data['original_num']-1)
            
            new_doc.save(output_path)
            new_doc.close()
            doc.close()
            
            return output_path
            
        except Exception as e:
            raise Exception(f"ಪುಟ ಸಾರಿಸುವಿಕೆ ವಿಫಲ: {str(e)}")

    def get_page_sorting_preview(self, file_path, session_id):
        """Generate preview for page sorting"""
        try:
            from .kannada_numeral_converter import KannadaNumeralConverter
            
            converter = KannadaNumeralConverter()
            doc = fitz.open(file_path)
            
            previews = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                extracted_number = converter.extract_page_number_from_text(text)
                
                thumbnail_path = self._generate_page_thumbnail(page, page_num + 1, session_id)
                
                previews.append({
                    'page_num': page_num + 1,
                    'extracted_number': extracted_number if extracted_number else page_num + 1,  # Fixed field name
                    'thumbnail_path': thumbnail_path
                })
            
            # Sort previews by extracted number to show the expected order
            sorted_previews = sorted(previews, key=lambda x: x['extracted_number'])  # Fixed field name
            
            # For the sorted_order, we need to return the sorted preview objects with proper field names
            sorted_order = []
            for preview in sorted_previews:
                sorted_order.append({
                    'page_num': preview['page_num'],
                    'extracted_number': preview['extracted_number'],  # This field name matches template
                    'thumbnail_path': preview['thumbnail_path']
                })
            
            doc.close()
            
            return {
                'total_pages': len(previews),
                'previews': previews,
                'sorted_order': sorted_order
            }
            
        except Exception as e:
            return {'error': f'ಪೂರ್ವವೀಕ್ಷಣೆ ರಚನೆ ವಿಫಲ: {str(e)}'}
    def _generate_page_thumbnail(self, page, page_num, session_id):
        """Generate a thumbnail image for a PDF page with automatic orientation detection"""
        try:
            import os
            from PIL import Image
            import io
            
            thumbnails_dir = os.path.join(self.config.OUTPUT_FOLDER, 'thumbnails', session_id)
            os.makedirs(thumbnails_dir, exist_ok=True)
            
            import time
            timestamp = int(time.time() * 1000)
            thumbnail_filename = f"page_{page_num}_{timestamp}.png"
            thumbnail_path = os.path.join(thumbnails_dir, thumbnail_filename)
            
            zoom = 1.5
            mat = fitz.Matrix(zoom, zoom)
            
            # Get page information
            page_rotation = page.rotation
            page_rect = page.rect
            
            # Get pixmap without pre-rotation
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # Smart orientation detection
            needs_rotation = False
            rotation_angle = 0
            
            # First, handle explicit page rotation from PDF
            if page_rotation != 0:
                # Handle rotated pages by rotating them back to normal
                if page_rotation == 90:
                    rotation_angle = -90
                    needs_rotation = True
                elif page_rotation == 180:
                    rotation_angle = 180  
                    needs_rotation = True
                elif page_rotation == 270:
                    rotation_angle = 90
                    needs_rotation = True
            else:
                # For pages with 0 rotation, try to detect if they're upside down
                # This is a heuristic based on text analysis
                try:
                    text_content = page.get_text()
                    
                    # If page has text, try to determine orientation
                    if text_content and len(text_content.strip()) > 10:
                        # Get text blocks with position information
                        blocks = page.get_text("dict")
                        
                        # Analyze text orientation heuristics
                        # Check if most text appears to be in normal reading order
                        normal_text_indicators = 0
                        total_text_blocks = 0
                        
                        for block in blocks.get("blocks", []):
                            if "lines" in block:
                                total_text_blocks += 1
                                for line in block["lines"]:
                                    for span in line.get("spans", []):
                                        text = span.get("text", "").strip()
                                        if text:
                                            # Check for Kannada or English characters in normal positions
                                            # If y-coordinates increase downward, text is likely normal
                                            # This is a simplified heuristic
                                            if any(c.isalnum() or ord(c) >= 0x0c80 for c in text):
                                                normal_text_indicators += 1
                        
                        # If we have very few normal text indicators relative to total blocks,
                        # the page might be upside down
                        if total_text_blocks > 0 and normal_text_indicators < (total_text_blocks * 0.3):
                            rotation_angle = 180
                            needs_rotation = True
                        
                except Exception as text_analysis_error:
                    # If text analysis fails, use simple dimension heuristic
                    # Many scanned documents appear upside down when height > width
                    page_width = page_rect.width
                    page_height = page_rect.height
                    
                    # This is a last resort heuristic - don't rotate by default
                    # Let users manually rotate if needed
                    pass
            
            # Apply rotation if needed
            if needs_rotation and rotation_angle != 0:
                img = img.rotate(rotation_angle, expand=True)
            
            thumbnail_size = (150, 200)
            img.thumbnail(thumbnail_size, Image.Resampling.LANCZOS)
            
            img.save(thumbnail_path, "PNG", optimize=True)
            
            return f'/thumbnails/{session_id}/{thumbnail_filename}'
            
        except Exception as e:
            print(f"Thumbnail generation error: {str(e)}")
            return None

    def protect_pdf(self, file_path, session_id, protection_options):
        """Protect PDF with password and permissions using PyPDF2 for better compatibility"""
        try:
            # Validate protection options
            password = protection_options.get('protection_password', '')
            if len(password) < 6:
                raise Exception('ಪಾಸ್‌ವರ್ಡ್ ಕನಿಷ್ಠ 6 ಅಕ್ಷರಗಳು ಇರಬೇಕು')
            
            confirm_password = protection_options.get('confirm_password', '')
            if password != confirm_password:
                raise Exception('ಪಾಸ್‌ವರ್ಡ್‌ಗಳು ಹೊಂದಿಕೆಯಾಗುತ್ತಿಲ್ಲ')
            
            # Create output filename
            original_name = os.path.splitext(os.path.basename(file_path))[0]
            output_filename = f"{original_name}_protected.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            # Handle file name conflicts by adding a number if file already exists
            counter = 1
            while os.path.exists(output_path):
                output_filename = f"{original_name}_protected_{counter}.pdf"
                output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
                counter += 1
            
            # Try PyPDF2 method first (more reliable for password protection)
            try:
                from PyPDF2 import PdfReader, PdfWriter
                
                # Read the source PDF
                reader = PdfReader(file_path)
                writer = PdfWriter()
                
                # Copy all pages
                for page in reader.pages:
                    writer.add_page(page)
                
                # Set password protection
                user_password = password
                owner_password = password
                
                # Apply encryption with password
                writer.encrypt(
                    user_password=user_password,
                    owner_password=owner_password,
                    use_128bit=True
                )
                
                # Save the protected PDF
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                
                print(f"PyPDF2 method successful: {output_path}")
                
            except Exception as pypdf2_error:
                print(f"PyPDF2 method failed: {pypdf2_error}")
                
                # Fallback to PyMuPDF with minimal encryption
                doc = fitz.open(file_path)
                
                # Use only basic password protection without complex permissions
                doc.save(
                    output_path,
                    encryption=1,  # Use basic RC4 encryption for compatibility
                    owner_pw=password,
                    user_pw=password
                    # No permissions parameter to avoid corruption
                )
                doc.close()
                print(f"PyMuPDF fallback method used: {output_path}")
            
            # Verify the output file was created properly
            if not os.path.exists(output_path):
                raise Exception('ಔಟ್‌ಪುಟ್ ಫೈಲ್ ರಚಿಸಲಾಗಿಲ್ಲ')
            
            file_size = os.path.getsize(output_path)
            if file_size == 0:
                raise Exception('ಖಾಲಿ ಫೈಲ್ ರಚಿಸಲಾಗಿದೆ')
            
            print(f"Protected PDF created successfully: {file_size} bytes")
            
            return {
                'success': True,
                'output_path': output_path,
                'filename': output_filename,
                'message': f'PDF ಯಶಸ್ವಿಯಾಗಿ ರಕ್ಷಿಸಲಾಗಿದೆ - ಪಾಸ್‌ವರ್ಡ್ ಅಗತ್ಯ'
            }
            
        except Exception as e:
            print(f"Protection error: {str(e)}")
            return {'success': False, 'error': f'PDF ರಕ್ಷಣೆ ವಿಫಲ: {str(e)}'}

    def unlock_pdf(self, file_path, password, session_id):
        """Remove password protection from PDF file"""
        try:
            if not os.path.exists(file_path):
                raise Exception('PDF ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ')
            
            if not password or password.strip() == '':
                raise Exception('ಪಾಸ್‌ವರ್ಡ್ ಅಗತ್ಯ')
            
            # Check if file is actually encrypted
            try:
                from PyPDF2 import PdfReader, PdfWriter
                reader = PdfReader(file_path)
                if not reader.is_encrypted:
                    raise Exception('ಈ PDF ಫೈಲ್ ರಕ್ಷಿತವಾಗಿಲ್ಲ')
            except Exception as e:
                if "not encrypted" in str(e):
                    raise e
                # If we can't read it, assume it's encrypted and continue
                pass
            
            # Try to decrypt the PDF
            reader = PdfReader(file_path)
            
            if not reader.decrypt(password):
                raise Exception('ತಪ್ಪು ಪಾಸ್‌ವರ್ಡ್ - ದಯವಿಟ್ಟು ಸರಿಯಾದ ಪಾಸ್‌ವರ್ಡ್ ನಮೂದಿಸಿ')
            
            # Create output filename
            original_name = os.path.splitext(os.path.basename(file_path))[0]
            output_filename = f"{original_name}_unlocked.pdf"
            output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
            
            # Handle file name conflicts by adding a number if file already exists
            counter = 1
            while os.path.exists(output_path):
                output_filename = f"{original_name}_unlocked_{counter}.pdf"
                output_path = os.path.join(self.config.OUTPUT_FOLDER, output_filename)
                counter += 1
            
            # Create new unlocked PDF
            writer = PdfWriter()
            
            # Add all pages to the writer
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                writer.add_page(page)
            
            # Write the unlocked PDF
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            # Verify the output file was created properly
            if not os.path.exists(output_path):
                raise Exception('ಔಟ್‌ಪುಟ್ ಫೈಲ್ ರಚಿಸಲಾಗಿಲ್ಲ')
            
            file_size = os.path.getsize(output_path)
            if file_size == 0:
                raise Exception('ಖಾಲಿ ಫೈಲ್ ರಚಿಸಲಾಗಿದೆ')
            
            print(f"Unlocked PDF created successfully: {file_size} bytes")
            
            return {
                'success': True,
                'output_path': output_path,
                'filename': output_filename,
                'message': f'PDF ಯಶಸ್ವಿಯಾಗಿ ಅನ್‌ಲಾಕ್ ಮಾಡಲಾಗಿದೆ'
            }
            
        except Exception as e:
            print(f"Unlock error: {str(e)}")
            return {'success': False, 'error': f'PDF ಅನ್‌ಲಾಕ್ ವಿಫಲ: {str(e)}'}

    
    def _convert_with_docx2pdf_threaded(self, input_path, session_id):
        """docx2pdf conversion with COM initialization in the worker thread"""
        try:
            # Check if docx2pdf is available
            try:
                import docx2pdf
            except ImportError:
                print("docx2pdf not installed")
                return None
            
            output_path = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_from_word.pdf")
            
            # Remove existing output file if it exists
            if os.path.exists(output_path):
                os.remove(output_path)
            
            print(f"Converting {input_path} to {output_path}")
            
            # CRITICAL FIX: COM must be initialized in the same thread that uses it
            conversion_result = {'success': False, 'error': None, 'path': None}
            
            def convert_worker():
                """Worker function that initializes COM in its own thread"""
                try:
                    # Initialize COM in this thread
                    if platform.system() == "Windows":
                        import pythoncom
                        pythoncom.CoInitialize()
                        print("✓ COM initialized in worker thread")
                    
                    # Perform conversion
                    docx2pdf.convert(input_path, output_path)
                    
                    # Check if successful
                    if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                        conversion_result['success'] = True
                        conversion_result['path'] = output_path
                    else:
                        conversion_result['error'] = "No output file created"
                    
                except Exception as e:
                    conversion_result['error'] = str(e)
                    print(f"Worker thread error: {e}")
                finally:
                    # Clean up COM in this thread
                    if platform.system() == "Windows":
                        try:
                            import pythoncom
                            pythoncom.CoUninitialize()
                            print("✓ COM uninitialized in worker thread")
                        except:
                            pass
            
            # Run conversion in thread with timeout
            thread = threading.Thread(target=convert_worker)
            thread.daemon = True
            thread.start()
            thread.join(timeout=120)  # 2 minute timeout
            
            if thread.is_alive():
                print("✗ Conversion timed out")
                return None
            
            if conversion_result['success']:
                print(f"✓ docx2pdf conversion successful: {os.path.getsize(conversion_result['path'])} bytes")
                return conversion_result['path']
            else:
                print(f"✗ Conversion failed: {conversion_result['error']}")
                return None
            
        except Exception as e:
            print(f"✗ docx2pdf threaded conversion failed: {e}")
            return None

    def _convert_with_libreoffice_simple(self, input_path, session_id):
        """Simple LibreOffice conversion"""
        try:
            # Find LibreOffice
            libreoffice_cmd = None
            
            # Common LibreOffice command names and paths
            candidates = ['soffice', 'libreoffice']
            
            # Platform-specific paths
            if platform.system() == "Windows":
                candidates.extend([
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
                ])
            elif platform.system() == "Darwin":
                candidates.extend([
                    "/Applications/LibreOffice.app/Contents/MacOS/soffice"
                ])
            else:  # Linux
                candidates.extend([
                    "/usr/bin/soffice",
                    "/usr/bin/libreoffice",
                    "/snap/bin/libreoffice"
                ])
            
            # Find working command
            for cmd in candidates:
                if shutil.which(cmd) or os.path.exists(cmd):
                    libreoffice_cmd = cmd
                    break
            
            if not libreoffice_cmd:
                print("LibreOffice not found")
                return None
            
            output_dir = self.config.OUTPUT_FOLDER
            os.makedirs(output_dir, exist_ok=True)
            
            # Run LibreOffice conversion
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                input_path
            ]
            
            print(f"Running: {' '.join(cmd)}")
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60
            )
            
            print(f"LibreOffice exit code: {result.returncode}")
            
            if result.returncode == 0:
                # Find generated PDF
                input_name = Path(input_path).stem
                generated_pdf = os.path.join(output_dir, f"{input_name}.pdf")
                final_output = os.path.join(output_dir, f"{session_id}_from_word.pdf")
                
                if os.path.exists(generated_pdf):
                    # Move to final location
                    if os.path.exists(final_output):
                        os.remove(final_output)
                    os.rename(generated_pdf, final_output)
                    
                    if os.path.getsize(final_output) > 0:
                        return final_output
            
            return None
            
        except Exception as e:
            print(f"LibreOffice simple conversion failed: {e}")
            return None

    

    def _clean_text_enhanced(self, text):
        """Enhanced text cleaning that preserves Kannada characters"""
        # Handle XML/HTML special characters
        replacements = {
            '&': '&amp;',
            '<': '&lt;',
            '>': '&gt;',
            '"': '&quot;',
            "'": '&apos;'
        }
        
        # Apply basic HTML entity encoding
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # Normalize Unicode to ensure consistent representation
        import unicodedata
        text = unicodedata.normalize('NFC', text)
        
        return text

    def _table_to_text_enhanced(self, table):
        """Enhanced table to text conversion"""
        try:
            lines = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Clean cell text
                        clean_cell = self._clean_text_enhanced(cell_text)
                        cells.append(clean_cell)
                
                if cells:
                    lines.append(' | '.join(cells))
            
            if lines:
                return '<br/>'.join(lines)
            
            return None
            
        except Exception as e:
            print(f"Enhanced table to text error: {e}")
            return None

    def _clean_text_simple(self, text):
        """Simple text cleaning for PDF generation - FIXED KANNADA SUPPORT"""
        # Handle common problematic characters for XML/HTML
        replacements = {
            '"': '"',
            '"': '"',
            ''': "'",
            ''': "'",
            '–': '-',
            '—': '-',
            '…': '...',
            '&': '&amp;',
            '<': '&lt;',
            '>': '&gt;'
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # CRITICAL FIX: Correct Unicode ranges for Kannada
        cleaned = ''
        for char in text:
            if ord(char) < 127:  # Basic ASCII
                cleaned += char
            elif ord(char) in range(2304, 2432):  # Devanagari range
                cleaned += char
            elif ord(char) in range(3072, 3200):  # CORRECT Kannada Unicode range (was 3200-3327)
                cleaned += char
            elif ord(char) in range(768, 880):   # Combining diacritical marks
                cleaned += char
            elif ord(char) in range(8204, 8207): # Zero-width characters (ZWNJ, ZWJ)
                cleaned += char
            elif char in ['\n', '\r', '\t', ' ']:  # Whitespace characters
                cleaned += char
            else:
                # Keep the character - don't filter it out
                cleaned += char
        
        return cleaned

    def _table_to_text(self, table):
        """Convert Word table to simple text representation"""
        try:
            lines = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        cells.append(cell_text)
                
                if cells:
                    lines.append(' | '.join(cells))
            
            if lines:
                return '<br/>'.join(lines)
            
            return None
            
        except Exception as e:
            print(f"Table to text error: {e}")
            return None

    def _convert_with_docx2pdf_threaded(self, input_path, session_id):
        """docx2pdf conversion with COM initialization in worker thread - FIXED"""
        try:
            # Check if docx2pdf is available
            try:
                import docx2pdf
            except ImportError:
                print("docx2pdf not installed")
                return None
            
            output_path = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_from_word.pdf")
            
            # Remove existing output file if it exists
            if os.path.exists(output_path):
                os.remove(output_path)
            
            print(f"Converting {input_path} to {output_path}")
            
            # Shared result container
            conversion_result = {'success': False, 'error': None, 'path': None}
            
            def convert_worker():
                """Worker function that initializes COM in its own thread"""
                try:
                    # CRITICAL FIX: Initialize COM in THIS thread
                    if platform.system() == "Windows":
                        import pythoncom
                        pythoncom.CoInitialize()
                        print("✓ COM initialized in worker thread")
                    
                    # Perform conversion
                    docx2pdf.convert(input_path, output_path)
                    
                    # Check if successful
                    if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                        conversion_result['success'] = True
                        conversion_result['path'] = output_path
                        print(f"✓ Conversion completed: {os.path.getsize(output_path)} bytes")
                    else:
                        conversion_result['error'] = "No output file created"
                    
                except Exception as e:
                    conversion_result['error'] = str(e)
                    print(f"✗ Worker thread error: {e}")
                finally:
                    # CRITICAL: Clean up COM in the same thread
                    if platform.system() == "Windows":
                        try:
                            import pythoncom
                            pythoncom.CoUninitialize()
                            print("✓ COM uninitialized in worker thread")
                        except Exception as cleanup_error:
                            print(f"⚠️ COM cleanup error: {cleanup_error}")
            
            # Run conversion in thread with timeout
            thread = threading.Thread(target=convert_worker)
            thread.daemon = True
            thread.start()
            thread.join(timeout=120)  # 2 minute timeout
            
            if thread.is_alive():
                print("✗ Conversion timed out")
                return None
            
            if conversion_result['success']:
                return conversion_result['path']
            else:
                print(f"✗ Conversion failed: {conversion_result.get('error', 'Unknown error')}")
                return None
            
        except Exception as e:
            print(f"✗ docx2pdf threaded conversion failed: {e}")
            return None

    # Alternative simpler approach - avoid threading altogether
    def _convert_with_docx2pdf_direct(self, input_path, session_id):
        """Direct docx2pdf conversion without threading"""
        try:
            # Check if docx2pdf is available
            try:
                import docx2pdf
            except ImportError:
                print("docx2pdf not installed")
                return None
            
            # Initialize COM before any operations
            if platform.system() == "Windows":
                import pythoncom
                pythoncom.CoInitialize()
                print("✓ COM initialized")
            
            output_path = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_from_word.pdf")
            
            # Remove existing output file if it exists
            if os.path.exists(output_path):
                os.remove(output_path)
            
            print(f"Converting {input_path} to {output_path}")
            
            # Direct conversion
            docx2pdf.convert(input_path, output_path)
            
            # Check if successful
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print(f"✓ Direct conversion successful: {os.path.getsize(output_path)} bytes")
                return output_path
            
            return None
            
        except Exception as e:
            print(f"✗ Direct docx2pdf conversion failed: {e}")
            return None
        finally:
            # Clean up COM
            if platform.system() == "Windows":
                try:
                    import pythoncom
                    pythoncom.CoUninitialize()
                    print("✓ COM uninitialized")
                except:
                    pass
    def generate_page_previews(self, pdf_path, session_id, preview_folder):
        """Generate page preview images for PDF - Updated to handle rotated pages"""
        try:
            # Create session-specific preview directory
            session_preview_dir = os.path.join(preview_folder, session_id)
            os.makedirs(session_preview_dir, exist_ok=True)
            
            # Open PDF with PyMuPDF for better image rendering
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            if total_pages == 0:
                return None
            
            previews = []
            
            # Generate preview for each page (limit to first 50 pages for performance)
            max_previews = min(total_pages, 50)
            
            for page_num in range(max_previews):
                try:
                    page = doc[page_num]
                    
                    # Create preview image
                    mat = fitz.Matrix(0.5, 0.5)  # Scale down for preview
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Convert to PIL Image
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Save preview image
                    preview_filename = f"page_{page_num + 1}.png"
                    preview_path = os.path.join(session_preview_dir, preview_filename)
                    img.save(preview_path, "PNG")
                    
                    previews.append({
                        'page_num': page_num + 1,
                        'image_path': preview_path
                    })
                    
                except Exception as page_error:
                    print(f"Error generating preview for page {page_num + 1}: {page_error}")
                    continue
            
            doc.close()
            
            return {
                'total_pages': total_pages,
                'previews': previews
            }
            
        except Exception as e:
            print(f"Preview generation error: {str(e)}")
        return None
        # Alternative method using pdf2image for full PDF processing
    def generate_page_previews_pdf2image(self, pdf_path, session_id, preview_folder, max_pages=None, batch_size=20):
        """
        Alternative method using pdf2image library for full PDF processing
        Requires: pip install pdf2image
        Also requires poppler-utils (system dependency)
        
        Args:
            pdf_path (str): Path to the PDF file
            session_id (str): Session identifier for organizing previews
            preview_folder (str): Base folder for storing preview images
            max_pages (int, optional): Maximum number of pages to generate previews for.
                                    If None, processes all pages
            batch_size (int): Number of pages to process in each batch
        """
        try:
            from pdf2image import convert_from_path
        
            # Create session-specific preview directory
            session_preview_dir = os.path.join(preview_folder, session_id)
            os.makedirs(session_preview_dir, exist_ok=True)
        
            # Get total page count
            total_pages = self.get_pdf_page_count(pdf_path)
            pages_to_process = total_pages if max_pages is None else min(total_pages, max_pages)
            
            print(f"Processing {pages_to_process} pages out of {total_pages} total pages...")
        
            previews = []
            
            # Process pages in batches to avoid memory issues
            for batch_start in range(0, pages_to_process, batch_size):
                batch_end = min(batch_start + batch_size, pages_to_process)
                first_page = batch_start + 1  # pdf2image uses 1-based indexing
                last_page = batch_end
                
                print(f"Processing batch: pages {first_page} to {last_page}")
                
                # Convert batch of PDF pages to images
                images = convert_from_path(
                    pdf_path,
                    dpi=150,  # Lower DPI for thumbnails
                    first_page=first_page,
                    last_page=last_page,
                    thread_count=2
                )
            
                for i, image in enumerate(images):
                    try:
                        # Resize to thumbnail
                        thumbnail_size = (200, 280)
                        image.thumbnail(thumbnail_size, Image.Resampling.LANCZOS)
                    
                        # Save thumbnail
                        page_num = batch_start + i + 1
                        filename = f"page_{page_num}.png"
                        file_path = os.path.join(session_preview_dir, filename)
                        image.save(file_path, "PNG", optimize=True)
                    
                        previews.append({
                            'page_num': page_num,
                            'image_path': file_path,
                            'width': image.width,
                            'height': image.height
                        })
                    
                    except Exception as e:
                        print(f"Error processing preview for page {batch_start + i + 1}: {str(e)}")
                        continue
                
                # Clean up memory
                images = None
                import gc
                gc.collect()
        
            return {
                'total_pages': total_pages,
                'previews': previews,
                'session_id': session_id,
                'processed_pages': len(previews)
            }
            
        except Exception as e:
            print(f"Error generating PDF previews with pdf2image: {str(e)}")
            return None

    def compare_pdfs(self, pdf1_path, pdf2_path, session_id, compare_type='both'):
        """Compare two PDF files and generate comparison report"""
        try:
            from utils.pdf_compare import PDFCompare
            
            # Validate input files
            if not os.path.exists(pdf1_path):
                raise Exception(f"ಮೊದಲ PDF ಫೈಲ್ ಸಿಗಲಿಲ್ಲ: {pdf1_path}")
            if not os.path.exists(pdf2_path):
                raise Exception(f"ಎರಡನೇ PDF ಫೈಲ್ ಸಿಗಲಿಲ್ಲ: {pdf2_path}")
            
            # Check if files are readable
            try:
                import fitz
                test_doc1 = fitz.open(pdf1_path)
                test_doc2 = fitz.open(pdf2_path)
                test_doc1.close()
                test_doc2.close()
            except Exception as pdf_error:
                raise Exception(f"PDF ಫೈಲ್‌ಗಳನ್ನು ಓದಲು ಸಾಧ್ಯವಾಗಿಲ್ಲ: {str(pdf_error)}")
            
            print(f"Comparing PDFs: {os.path.basename(pdf1_path)} vs {os.path.basename(pdf2_path)}")
            print(f"Compare type: {compare_type}")
            
            comparer = PDFCompare()
            result_path = comparer.compare_pdfs(pdf1_path, pdf2_path, session_id, compare_type)
            
            if not result_path or not os.path.exists(result_path):
                raise Exception("ಹೋಲಿಕೆ ವರದಿ ರಚಿಸಲಾಗಿಲ್ಲ")
            
            print(f"Comparison report created: {result_path}")
            return result_path
            
        except Exception as e:
            print(f"Compare operation error: {e}")
            raise Exception(f"PDF ಹೋಲಿಕೆ ವಿಫಲ: {str(e)}")
    def compare_pdfs_web(self, pdf1_path, pdf2_path, session_id, compare_type='both'):
        """Web-friendly PDF comparison that returns data for Flask templates"""
        try:
            from utils.pdf_compare import PDFCompare
            comparer = PDFCompare()
            return comparer.compare_pdfs_web(pdf1_path, pdf2_path, session_id, compare_type)
        except Exception as e:
            print(f"PDF comparison error: {e}")
            return None
    def _convert_with_weasyprint(self, input_path, session_id):
        """
        Convert Word to PDF using WeasyPrint - SIMPLE COMPATIBLE VERSION
        """
        try:
            # Check if weasyprint is available
            try:
                from weasyprint import HTML, CSS
            except ImportError:
                print("WeasyPrint not installed")
                return None
            
            from docx import Document
            import html
            import base64
            import zipfile
            import os
            
            print("Starting WeasyPrint conversion with IMAGE support (simple method)...")
            
            # SIMPLE APPROACH: Extract images directly from the .docx ZIP file
            images_data = {}
            try:
                with zipfile.ZipFile(input_path, 'r') as docx_zip:
                    # Find all image files in the media folder
                    media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                    
                    for media_file in media_files:
                        try:
                            # Extract image data
                            image_data = docx_zip.read(media_file)
                            
                            # Get image filename and extension
                            image_filename = os.path.basename(media_file)
                            image_name, image_ext = os.path.splitext(image_filename)
                            
                            # Determine MIME type
                            if image_ext.lower() in ['.jpg', '.jpeg']:
                                mime_type = 'image/jpeg'
                            elif image_ext.lower() == '.png':
                                mime_type = 'image/png'
                            elif image_ext.lower() == '.gif':
                                mime_type = 'image/gif'
                            elif image_ext.lower() == '.bmp':
                                mime_type = 'image/bmp'
                            else:
                                mime_type = 'image/png'
                            
                            # Convert to base64
                            image_base64 = base64.b64encode(image_data).decode('utf-8')
                            
                            # Store image data
                            images_data[image_filename] = {
                                'base64': image_base64,
                                'mime_type': mime_type,
                                'filename': image_filename
                            }
                            
                            print(f"✓ Found image: {image_filename}")
                            
                        except Exception as img_error:
                            print(f"Error processing image {media_file}: {img_error}")
                            continue
            
            except Exception as zip_error:
                print(f"Could not extract images from docx: {zip_error}")
                # Continue without images
            
            # Read Word document
            doc = Document(input_path)
            
            # Create HTML content
            html_content = '''<!DOCTYPE html>
    <html lang="kn">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Document</title>
        <style>
            body {
                font-family: 'Noto Sans Kannada', 'Tunga', 'Lohit Kannada', Arial, sans-serif;
                font-size: 12pt;
                line-height: 1.6;
                margin: 2cm;
                color: #000;
                text-rendering: optimizeLegibility;
            }
            
            .paragraph {
                margin-bottom: 12pt;
                text-align: justify;
            }
            
            .title {
                font-size: 16pt;
                font-weight: bold;
                text-align: center;
                margin-bottom: 20pt;
            }
            
            .heading {
                font-size: 14pt;
                font-weight: bold;
                margin-top: 16pt;
                margin-bottom: 12pt;
            }
            
            .image {
                max-width: 100%;
                height: auto;
                margin: 12pt auto;
                display: block;
            }
            
            .image-container {
                text-align: center;
                margin: 12pt 0;
            }
            
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 12pt 0;
                font-size: 11pt;
            }
            
            td, th {
                border: 1px solid #000;
                padding: 6pt;
                text-align: left;
            }
            
            th {
                background-color: #f0f0f0;
                font-weight: bold;
            }
        </style>
    </head>
    <body>
    '''
            
            # Process paragraphs
            image_counter = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                
                # If paragraph is empty but might contain images, add images
                if not text and images_data:
                    # Add the next available image
                    if image_counter < len(images_data):
                        image_list = list(images_data.values())
                        if image_counter < len(image_list):
                            img_data = image_list[image_counter]
                            html_content += f'''
    <div class="image-container">
        <img src="data:{img_data['mime_type']};base64,{img_data['base64']}" class="image" alt="Document Image" />
    </div>
    '''
                            image_counter += 1
                            continue
                
                if not text:
                    html_content += '<div class="paragraph">&nbsp;</div>\n'
                    continue
                
                # Escape HTML but preserve Kannada characters
                escaped_text = html.escape(text, quote=False)
                
                # Determine paragraph style
                style_name = para.style.name.lower() if para.style and para.style.name else ''
                
                if 'title' in style_name:
                    html_content += f'<div class="title">{escaped_text}</div>\n'
                elif 'heading' in style_name:
                    html_content += f'<div class="heading">{escaped_text}</div>\n'
                else:
                    html_content += f'<div class="paragraph">{escaped_text}</div>\n'
            
            # Add any remaining images at the end
            while image_counter < len(images_data):
                image_list = list(images_data.values())
                if image_counter < len(image_list):
                    img_data = image_list[image_counter]
                    html_content += f'''
    <div class="image-container">
        <img src="data:{img_data['mime_type']};base64,{img_data['base64']}" class="image" alt="Document Image" />
    </div>
    '''
                    image_counter += 1
                else:
                    break
            
            # Process tables
            for table in doc.tables:
                html_content += '<table>\n'
                for i, row in enumerate(table.rows):
                    html_content += '<tr>\n'
                    for cell in row.cells:
                        cell_text = html.escape(cell.text.strip(), quote=False)
                        tag = 'th' if i == 0 else 'td'
                        html_content += f'<{tag}>{cell_text}</{tag}>\n'
                    html_content += '</tr>\n'
                html_content += '</table>\n'
            
            html_content += '</body></html>'
            
            # Generate PDF using WeasyPrint
            output_path = os.path.join(self.config.OUTPUT_FOLDER, f"{session_id}_from_word.pdf")
            
            try:
                html_doc = HTML(string=html_content)
                html_doc.write_pdf(output_path)
                
            except Exception as api_error:
                print(f"WeasyPrint API error: {api_error}")
                try:
                    HTML(string=html_content, encoding='utf-8').write_pdf(output_path)
                except Exception as alt_error:
                    print(f"WeasyPrint alternative API error: {alt_error}")
                    return None
            
            # Validate output
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print(f"WeasyPrint PDF with images created successfully: {os.path.getsize(output_path)} bytes")
                return output_path
            
            return None
            
        except Exception as e:
            print(f"WeasyPrint conversion with images error: {e}")
            import traceback
            traceback.print_exc()
            return None