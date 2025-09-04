#!/usr/bin/env python3
"""
PDF Text Extractor Script
This script extracts text from PDF files and saves them to txt files.
Supports both text-based PDFs and image-based PDFs (using OCR).
"""

import os
import sys
import argparse
from datetime import datetime

# PDF text extraction libraries
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# OCR libraries for image-based PDFs
try:
    import pytesseract
    from PIL import Image
    import pdf2image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Word document library
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def check_dependencies():
    """Check if required dependencies are installed."""
    print("🔍 Checking dependencies...")
    
    missing_deps = []
    
    if not PYPDF2_AVAILABLE and not PDFPLUMBER_AVAILABLE:
        missing_deps.append("PyPDF2 or pdfplumber")
    
    if PYPDF2_AVAILABLE:
        print("✅ PyPDF2 is available")
    elif PDFPLUMBER_AVAILABLE:
        print("✅ pdfplumber is available")
    
    if OCR_AVAILABLE:
        print("✅ OCR libraries (pytesseract, PIL, pdf2image) are available")
        
        # Check Tesseract installation
        try:
            tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            if os.path.exists(tesseract_path):
                print("✅ Tesseract OCR is properly installed")
                
                # Check Kannada language support
                tessdata_path = r'C:\Program Files\Tesseract-OCR\tessdata\kan.traineddata'
                if os.path.exists(tessdata_path):
                    print("✅ Kannada language pack is available")
                else:
                    print("⚠️  Kannada language pack not found")
            else:
                print("⚠️  Tesseract not found at expected location")
        except Exception as e:
            print(f"⚠️  Tesseract check failed: {e}")
            
        # Check Poppler installation
        poppler_path = r'C:\poppler\poppler-24.08.0\Library\bin\pdftoppm.exe'
        if os.path.exists(poppler_path):
            print("✅ Poppler is properly installed")
        else:
            print("⚠️  Poppler not found at expected location")
            
    else:
        print("⚠️  OCR libraries not available (for image-based PDFs)")
        missing_deps.append("pytesseract, pillow, pdf2image (for OCR)")
    
    if DOCX_AVAILABLE:
        print("✅ python-docx is available")
    else:
        print("⚠️  python-docx not available (for DOCX output)")
        missing_deps.append("python-docx (for DOCX output)")
    
    if missing_deps:
        print("\n❌ Missing dependencies:")
        for dep in missing_deps:
            print(f"   - {dep}")
        print("\nInstall with:")
        print("pip install PyPDF2 pdfplumber pytesseract pillow pdf2image python-docx")
        return False
    
    return True


def extract_text_pypdf2(pdf_path):
    """
    Extract text from PDF using PyPDF2.
    
    Args:
        pdf_path (str): Path to PDF file
    
    Returns:
        str: Extracted text
    """
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            print(f"📄 PDF has {len(pdf_reader.pages)} pages")
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                print(f"📖 Processing page {page_num}...")
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += f"\n--- Page {page_num} ---\n"
                    text += page_text + "\n"
            
            return text
    except Exception as e:
        return f"❌ Error extracting text with PyPDF2: {str(e)}"


def extract_text_pdfplumber(pdf_path):
    """
    Extract text from PDF using pdfplumber (better for complex layouts).
    
    Args:
        pdf_path (str): Path to PDF file
    
    Returns:
        str: Extracted text
    """
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            
            print(f"📄 PDF has {len(pdf.pages)} pages")
            
            for page_num, page in enumerate(pdf.pages, 1):
                print(f"📖 Processing page {page_num}...")
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += f"\n--- Page {page_num} ---\n"
                    text += page_text + "\n"
            
            return text
    except Exception as e:
        return f"❌ Error extracting text with pdfplumber: {str(e)}"


def extract_text_ocr(pdf_path, language='kan'):
    """
    High-quality Kannada OCR text extraction from PDF using Tesseract with Poppler.
    Optimized specifically for Kannada text recognition.
    
    Args:
        pdf_path (str): Path to PDF file
        language (str): OCR language code ('kan' for Kannada, 'kan+eng' for mixed)
    
    Returns:
        str: Extracted Kannada text
    """
    if not OCR_AVAILABLE:
        return "❌ OCR libraries not available. Install pytesseract, pillow, and pdf2image."
    
    # Configure Tesseract path for production use
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    
    try:
        # Convert PDF pages to high-quality images using Poppler
        print("🔄 Converting PDF to high-quality images for Kannada OCR...")
        
        # Use optimal DPI for Kannada text recognition
        images = pdf2image.convert_from_path(
            pdf_path,
            dpi=400,  # Higher DPI for better Kannada character recognition
            fmt='PNG',  # PNG for lossless quality
            thread_count=4,  # Use multiple threads for faster conversion
            poppler_path=r"C:\poppler\poppler-24.08.0\Library\bin"  # Explicit Poppler path
        )
        
        text = ""
        print(f"📄 PDF converted to {len(images)} high-quality image(s)")
        
        # Optimized Tesseract configuration for Kannada
        # PSM 6: Uniform block of text (best for Kannada documents)
        # OEM 3: Default LSTM OCR Engine Mode
        # Force Kannada-only recognition for best results
        kannada_config = r'--oem 3 --psm 6 -c tessedit_char_blacklist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz'
        
        print(f"🔧 Using Kannada-optimized OCR with language: {language}")
        
        for page_num, image in enumerate(images, 1):
            print(f"🔍 OCR processing page {page_num} for Kannada text...")
            
            # Specialized image preprocessing for Kannada text
            # Convert to grayscale for better Kannada character recognition
            from PIL import ImageEnhance, ImageFilter, ImageOps
            
            # Convert to grayscale first (better for Kannada OCR)
            image_gray = ImageOps.grayscale(image)
            
            # Convert back to RGB for compatibility
            image = image_gray.convert('RGB')
            
            # Optimize contrast specifically for Kannada text
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.2)  # Higher contrast for Kannada
            
            # Increase sharpness for better character edges
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(2.0)
            
            # Apply gentle smoothing to reduce noise while preserving character details
            image = image.filter(ImageFilter.SMOOTH_MORE)
            
            # FIRST: Try pure Kannada OCR
            print(f"   Attempting pure Kannada OCR...")
            page_text = pytesseract.image_to_string(
                image, 
                lang='kan',  # Pure Kannada only
                config=kannada_config
            )
            
            # If pure Kannada doesn't work well, try with English fallback
            if not page_text.strip() or len(page_text.strip()) < 10:
                print(f"   Pure Kannada OCR yielded little text, trying mixed mode...")
                page_text = pytesseract.image_to_string(
                    image, 
                    lang='kan+eng',  # Mixed Kannada and English
                    config=r'--oem 3 --psm 6'
                )
            
            # Clean and validate the extracted text
            if page_text.strip():
                # Remove obvious OCR errors and clean text
                cleaned_text = page_text.strip()
                
                # Check if we got meaningful Kannada text
                kannada_chars = sum(1 for char in cleaned_text if '\u0C80' <= char <= '\u0CFF')
                total_chars = len([c for c in cleaned_text if c.isalnum()])
                
                if total_chars > 0:
                    kannada_percentage = (kannada_chars / total_chars) * 100
                    print(f"   Extracted {len(cleaned_text)} characters ({kannada_percentage:.1f}% Kannada) from page {page_num}")
                else:
                    print(f"   Extracted {len(cleaned_text)} characters from page {page_num}")
                
                text += f"\n--- Page {page_num} (Kannada OCR) ---\n"
                text += cleaned_text + "\n"
        
        return text
    except Exception as e:
        return f"❌ Error extracting text with OCR: {str(e)}"


def extract_text_from_pdf(pdf_path, method='auto', ocr_language='kan'):
    """
    Extract text from PDF using specified method.
    
    Args:
        pdf_path (str): Path to PDF file
        method (str): Extraction method ('auto', 'pypdf2', 'pdfplumber', 'ocr')
        ocr_language (str): Language for OCR
    
    Returns:
        str: Extracted text
    """
    if not os.path.exists(pdf_path):
        return f"❌ PDF file '{pdf_path}' not found."
    
    print(f"📂 Processing PDF: {pdf_path}")
    
    if method == 'pypdf2' and PYPDF2_AVAILABLE:
        return extract_text_pypdf2(pdf_path)
    elif method == 'pdfplumber' and PDFPLUMBER_AVAILABLE:
        return extract_text_pdfplumber(pdf_path)
    elif method == 'ocr':
        return extract_text_ocr(pdf_path, ocr_language)
    elif method == 'auto':
        # Try pdfplumber first, then PyPDF2, then OCR
        if PDFPLUMBER_AVAILABLE:
            print("🔄 Trying pdfplumber extraction...")
            text = extract_text_pdfplumber(pdf_path)
            if text and not text.startswith("❌") and text.strip():
                return text
        
        if PYPDF2_AVAILABLE:
            print("🔄 Trying PyPDF2 extraction...")
            text = extract_text_pypdf2(pdf_path)
            if text and not text.startswith("❌") and text.strip():
                return text
        
        if OCR_AVAILABLE:
            print("🔄 Trying OCR extraction...")
            return extract_text_ocr(pdf_path, ocr_language)
        
        return "❌ No suitable extraction method available."
    else:
        return f"❌ Method '{method}' not available or not supported."


def save_text_to_file(text, pdf_path, output_dir=None):
    """
    Save extracted text to a .txt file.
    
    Args:
        text (str): Extracted text to save
        pdf_path (str): Path to the original PDF
        output_dir (str): Directory to save the text file (optional)
    
    Returns:
        str: Path to the saved text file
    """
    # Generate output filename based on PDF name
    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_filename = f"{pdf_name}_extracted_text.txt"
    
    # Determine output directory
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)
    else:
        # Save in the same directory as the PDF
        pdf_dir = os.path.dirname(pdf_path) or '.'
        output_path = os.path.join(pdf_dir, output_filename)
    
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            # Write header with metadata
            f.write(f"Extracted Text from: {os.path.basename(pdf_path)}\n")
            f.write(f"Processing Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Character Count: {len(text)}\n")
            f.write(f"Word Count: {len(text.split())}\n")
            f.write(f"Line Count: {len(text.splitlines())}\n")
            f.write("=" * 50 + "\n\n")
            
            # Write the extracted text
            f.write(text)
        
        return output_path
    except Exception as e:
        print(f"❌ Error saving text file: {str(e)}")
        return None


def save_text_to_docx(text, output_path):
    """Save extracted text to a DOCX file."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not available. Install with 'pip install python-docx'")
    # Ensure directory exists
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc = Document()
    # Split by lines to preserve basic structure
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(output_path)
    return output_path


def convert_pdf_to_docx(pdf_path, session_id, output_dir=None, method='ocr', ocr_language='kan'):
    """
    Convert a PDF to a DOCX by extracting text (auto/ocr) and writing to a Word document.
    Returns the path to the generated .docx file.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    # Extract text using the chosen method
    extracted_text = extract_text_from_pdf(pdf_path, method=method, ocr_language=ocr_language)
    if not extracted_text or extracted_text.startswith("❌"):
        raise RuntimeError(f"Text extraction failed: {extracted_text}")

    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_dir_final = output_dir or (os.path.dirname(pdf_path) or '.')
    output_filename = f"{session_id}_{base_name}.docx"
    output_path = os.path.join(output_dir_final, output_filename)

    # Save to DOCX
    save_text_to_docx(extracted_text, output_path)

    if not os.path.exists(output_path):
        raise RuntimeError("DOCX file was not created")

    return output_path


def convert_pdf_to_docx_with_images(
    pdf_path,
    session_id,
    output_dir=None,
    text_method='ocr',
    ocr_language='kan',
    include_images=False,
    image_dpi=200,
    image_max_width_inches=6.0,
    
    
):
    
   
    """
    Convert a PDF to a DOCX file.

    - Extracts text using the requested method (default OCR) and writes it to the DOCX.
    - Optionally embeds page images rendered from the PDF into the DOCX for visual reference.

    Args:
        pdf_path (str): Path to the PDF file.
        session_id (str): Session identifier used in the output filename.
        output_dir (str | None): Directory to write the DOCX into. Defaults to PDF directory.
        text_method (str): 'auto' | 'pypdf2' | 'pdfplumber' | 'ocr'. Defaults to 'ocr'.
        ocr_language (str): Tesseract language codes, e.g. 'kan+eng'. Defaults to 'kan'.
        include_images (bool): When True, embeds page images. Defaults to False.
        image_dpi (int): Rendering DPI for page images. Defaults to 200.
        image_max_width_inches (float): Max width for images in DOCX. Defaults to 6.0 inches.

    Returns:
        str: Path to the generated DOCX file.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not available. Install with 'pip install python-docx'")

    # Determine output path
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_dir_final = output_dir or (os.path.dirname(pdf_path) or '.')
    os.makedirs(output_dir_final, exist_ok=True)
    output_filename = f"{session_id}_{base_name}.docx"
    output_path = os.path.join(output_dir_final, output_filename)

    # Create document
    doc = Document()

    # Extract text and images together in their natural positions
    if include_images:
        print(f"🚀 EXTRACTING TEXT AND REAL IMAGES ONLY...")
        
        try:
            import io as _io
            from docx.shared import Inches
            import fitz  # PyMuPDF
            
            pdf_doc = fitz.open(pdf_path)
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc.load_page(page_num)
                
                # Add page heading
                doc.add_heading(f"Page {page_num + 1}", level=2)
                
                # Get all text blocks with their positions
                text_blocks = page.get_text("dict")
                
                # Get all ACTUAL images (not text rendered as images)
                image_list = page.get_images(full=True)
                
                # Create a list of all content items (text blocks + real images) with positions
                content_items = []
                
                # Add text blocks
                for block in text_blocks["blocks"]:
                    if "lines" in block:  # Text block
                        y_pos = block["bbox"][1]  # Top Y coordinate
                        
                        # FIXED: Properly combine text spans within lines
                        block_text = ""
                        for line in block["lines"]:
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                            
                            # Add the complete line with proper spacing
                            if line_text.strip():
                                block_text += line_text.strip() + " "
                        
                        # Clean up the complete block text
                        block_text = block_text.strip()
                        
                        if block_text:
                            content_items.append({
                                "type": "text",
                                "y_pos": y_pos,
                                "content": block_text
                            })
                
                # Add ONLY actual images (filter out text-based images)
                actual_images_found = 0
                for img_index, img in enumerate(image_list):
                    try:
                        # Get image position
                        img_rect = page.get_image_rects(img[0])
                        if img_rect:
                            y_pos = img_rect[0].y0  # Top Y coordinate of image
                            
                            # Extract the actual image
                            xref = img[0]
                            pix = fitz.Pixmap(pdf_doc, xref)
                            
                            # IMPROVED filtering: Include maps, diagrams, and actual pictures
                            # Skip very small images (likely icons/bullets) but include larger diagrams/maps
                            is_actual_image = (
                                pix.width >= 30 and pix.height >= 30 and  # Allow smaller diagrams/maps
                                pix.width <= 3000 and pix.height <= 3000 and  # Allow larger diagrams
                                (pix.width * pix.height) >= 1500  # Lower minimum area for diagrams/maps
                            )
                            
                            if is_actual_image:
                                # Handle color space conversion
                                if pix.n - pix.alpha < 4:  # Not CMYK
                                    if pix.n != 3:
                                        try:
                                            pix = fitz.Pixmap(fitz.csRGB, pix)
                                        except:
                                            pass
                                    
                                    img_data = pix.pil_tobytes(format="PNG")
                                    
                                    content_items.append({
                                        "type": "image",
                                        "y_pos": y_pos,
                                        "content": img_data,
                                        "width": pix.width,
                                        "height": pix.height
                                    })
                                    
                                    actual_images_found += 1
                                    print(f"      ✅ Found actual image: {pix.width}x{pix.height}px")
                            else:
                                print(f"      ⚠️  Skipped (likely text/icon): {pix.width}x{pix.height}px")
                                    
                            pix = None
                            
                    except Exception as img_err:
                        print(f"      ❌ Error processing image {img_index + 1}: {img_err}")
                
                # FOR HAND-DRAWN CONTENT: Check if page has vector graphics/drawings
                # If no embedded images found but page has drawings, render the page
                page_has_drawings = len(page.get_drawings()) > 0
                
                if actual_images_found == 0 and page_has_drawings:
                    print(f"      🖊️  Detected hand-drawn content, rendering page...")
                    try:
                        # Render the entire page to capture hand-drawn elements
                        mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for quality
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.pil_tobytes(format="PNG")
                        
                        # Add the rendered page as an image
                        content_items.append({
                            "type": "image",
                            "y_pos": 0,  # Place at top since it's the whole page
                            "content": img_data,
                            "width": pix.width,
                            "height": pix.height
                        })
                        
                        actual_images_found += 1
                        print(f"      ✅ Page rendered to capture drawings: {pix.width}x{pix.height}px")
                        pix = None
                        
                    except Exception as render_err:
                        print(f"      ❌ Failed to render page: {render_err}")
                
                # Sort all content by Y position (top to bottom)
                content_items.sort(key=lambda x: x["y_pos"])
                
                # Add content to document in correct order
                text_blocks_added = 0
                images_added = 0
                
                for item in content_items:
                    if item["type"] == "text":
                        # FIXED: Add as single paragraph with proper text flow
                        if item["content"].strip():
                            doc.add_paragraph(item["content"])
                            text_blocks_added += 1
                    
                    elif item["type"] == "image":
                        # Add actual image inline
                        buf = _io.BytesIO(item["content"])
                        
                        # Calculate appropriate size
                        aspect_ratio = item["width"] / item["height"]
                        if aspect_ratio > 1:  # Wider than tall
                            width = min(image_max_width_inches, 5.0)
                        else:  # Taller than wide
                            width = min(3.0, image_max_width_inches)
                        
                        doc.add_picture(buf, width=Inches(width))
                        images_added += 1
                        print(f"   🖼️  Image added inline ({item['width']}x{item['height']}px)")
                
                if text_blocks_added == 0 and images_added == 0:
                    doc.add_paragraph("(No extractable content on this page)")
                
                # Add space between pages
                doc.add_paragraph("")
                print(f"   ✅ Page {page_num + 1}: {text_blocks_added} text blocks, {images_added} actual images")
            
            pdf_doc.close()
            print(f"🎉 All pages processed - only real images included!")
                
        except Exception as major_err:
            print(f"❌ CRITICAL: Content extraction failed: {major_err}")
            doc.add_paragraph(f"ERROR: Content extraction failed - {major_err}")
            
            # Fallback to simple text extraction
            try:
                extracted_text = extract_text_from_pdf(pdf_path, method=text_method, ocr_language=ocr_language)
                if extracted_text and not str(extracted_text).startswith("❌"):
                    for line in str(extracted_text).splitlines():
                        doc.add_paragraph(line)
                else:
                    doc.add_paragraph(str(extracted_text) if extracted_text else "No text extracted.")
            except Exception as e:
                doc.add_paragraph(f"❌ Text extraction error: {e}")
        
        print(f"✅ TEXT AND REAL IMAGE EXTRACTION COMPLETE!")
    else:
        # Text-only mode
        print(f"ℹ️  Text-only mode (images disabled)")
        try:
            extracted_text = extract_text_from_pdf(pdf_path, method=text_method, ocr_language=ocr_language)
        except Exception as e:
            extracted_text = f"❌ Text extraction error: {e}"

        if extracted_text and not str(extracted_text).startswith("❌"):
            for line in str(extracted_text).splitlines():
                doc.add_paragraph(line)
        else:
            doc.add_paragraph(str(extracted_text) if extracted_text else "No text extracted.")

    # Save the DOCX
    doc.save(output_path)

    if not os.path.exists(output_path):
        raise RuntimeError("DOCX file was not created")

    return output_path


def display_results(text, pdf_path, save_output=False, output_dir=None):
    """
    Display the extraction results and optionally save to file.
    
    Args:
        text (str): Extracted text
        pdf_path (str): Path to the processed PDF
        save_output (bool): Whether to save text to file
        output_dir (str): Directory to save the text file
    """
    print("\n" + "="*60)
    print(f"📋 EXTRACTION RESULTS FOR: {os.path.basename(pdf_path)}")
    print("="*60)
    
    if text.strip() and not text.startswith("❌"):
        print("📝 Extracted Text Preview (first 500 characters):")
        print("-" * 40)
        preview = text[:500] + "..." if len(text) > 500 else text
        print(preview)
        print("-" * 40)
        print(f"📊 Character count: {len(text)}")
        print(f"📊 Word count: {len(text.split())}")
        print(f"📊 Line count: {len(text.splitlines())}")
        
        # Save to file if requested
        if save_output:
            output_path = save_text_to_file(text, pdf_path, output_dir)
            if output_path:
                print(f"💾 Text saved to: {output_path}")
    else:
        print("⚠️  No text was extracted from the PDF or an error occurred.")
        print("💡 Tips:")
        print("   - Try using --method ocr for image-based PDFs")
        print("   - Ensure the PDF is not password protected")
        print("   - Check if the PDF contains actual text or just images")


def process_multiple_pdfs(folder_path, method='auto', ocr_language='kan', save_output=False, output_dir=None):
    """
    Process all PDF files in a folder.
    
    Args:
        folder_path (str): Path to folder containing PDFs
        method (str): Extraction method
        ocr_language (str): Language for OCR
        save_output (bool): Whether to save text to files
        output_dir (str): Directory to save the text files
    """
    if not os.path.exists(folder_path):
        print(f"❌ Folder '{folder_path}' not found.")
        return
    
    pdf_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.pdf')]
    
    if not pdf_files:
        print(f"❌ No PDF files found in '{folder_path}'")
        return
    
    print(f"📁 Found {len(pdf_files)} PDF file(s) in '{folder_path}")
    
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n🔄 Processing PDF {i}/{len(pdf_files)}: {pdf_file}")
        pdf_path = os.path.join(folder_path, pdf_file)
        
        extracted_text = extract_text_from_pdf(pdf_path, method, ocr_language)
        display_results(extracted_text, pdf_path, save_output, output_dir)


def main():
    """Main function to handle command line arguments and execute PDF text extraction."""
    parser = argparse.ArgumentParser(
        description="Extract text from PDF files and save to txt files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python pdf_text_extractor.py document.pdf                    # Extract text from PDF
  python pdf_text_extractor.py document.pdf --save             # Extract and save to txt
  python pdf_text_extractor.py --folder ./pdfs --save          # Process folder and save
  python pdf_text_extractor.py document.pdf --method ocr       # Use OCR for image-based PDF
  python pdf_text_extractor.py document.pdf --save --output ./texts  # Save to specific directory
        """
    )
    
    parser.add_argument('path', nargs='?', help='Path to PDF file or folder')
    parser.add_argument('--folder', action='store_true',
                       help='Process all PDF files in the specified folder')
    parser.add_argument('--method', choices=['auto', 'pypdf2', 'pdfplumber', 'ocr'], 
                       default='auto', help='Text extraction method (default: auto)')
    parser.add_argument('--ocr-language', default='kan',
                       help='Language for OCR (default: kan for Kannada)')
    parser.add_argument('--save', action='store_true',
                       help='Save extracted text to .txt files')
    parser.add_argument('--output', type=str, metavar='DIR',
                       help='Directory to save text files (default: same as PDF location)')
    parser.add_argument('--check', action='store_true',
                       help='Check available dependencies')
    
    args = parser.parse_args()
    
    print("🚀 PDF Text Extractor")
    print("=" * 30)
    
    # Check dependencies if requested
    if args.check:
        check_dependencies()
        return
    
    # Check if path is provided
    if not args.path:
        print("❌ Please provide a PDF path or folder path.")
        print("Use --help for usage information.")
        return
    
    # Check dependencies
    if not check_dependencies():
        return
    
    # Validate output directory if specified
    if args.output and not args.save:
        print("⚠️  --output can only be used with --save option")
        return
    
    print()  # Add spacing
    
    # Process folder or single PDF
    if args.folder:
        process_multiple_pdfs(args.path, args.method, args.ocr_language, args.save, args.output)
    else:
        # Process single PDF
        extracted_text = extract_text_from_pdf(args.path, args.method, args.ocr_language)
        display_results(extracted_text, args.path, args.save, args.output)
    
    print("\n✅ Processing completed!")


if __name__ == "__main__":
    main()