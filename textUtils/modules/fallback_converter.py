from typing import Optional

import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)

class FallbackPDFConverter:
    """Simple PDF to Word converter that doesn't depend on poppler"""
    
    def __init__(self, debug_mode=False):
        self.debug_mode = debug_mode
        
    def convert_pdf_to_word_simple(
        self, 
        input_pdf_path: str, 
        output_docx_path: str,
        output_txt_path: Optional[str] = None,
        title: Optional[str] = None,
        author: Optional[str] = None
    ):
        """
        Simple PDF to Word conversion using PyMuPDF only
        Returns: (docx_path, txt_path, None)
        """
        try:
            # Open PDF
            doc = fitz.open(input_pdf_path)
            
            # Create Word document
            word_doc = Document()
            
            # Add title
            if title:
                title_para = word_doc.add_heading(title, 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                title_para = word_doc.add_heading('PDF ನಿಂದ ಪರಿವರ್ತಿತ ದಾಖಲೆ', 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            full_text = []
            
            # Extract text from each page
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text
                text = page.get_text()
                
                if text.strip():  # Only add non-empty pages
                    # Add page header
                    page_header = word_doc.add_heading(f'ಪುಟ {page_num + 1}', level=2)
                    
                    # Split text into paragraphs and add them
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        clean_para = para.strip()
                        if clean_para:
                            word_doc.add_paragraph(clean_para)
                            full_text.append(clean_para)
                    
                    # Add page break except for last page
                    if page_num < len(doc) - 1:
                        word_doc.add_page_break()
                
                # Also try to extract images
                try:
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        # Get image
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("png")
                            
                            # Save temporary image
                            temp_img_path = f"temp_img_{page_num}_{img_index}.png"
                            with open(temp_img_path, "wb") as f:
                                f.write(img_data)
                            
                            # Add image to Word document
                            try:
                                word_doc.add_picture(temp_img_path, width=Inches(6))
                                # Clean up temp file
                                os.remove(temp_img_path)
                            except Exception as img_err:
                                logger.warning(f"Could not add image: {img_err}")
                                if os.path.exists(temp_img_path):
                                    os.remove(temp_img_path)
                        
                        pix = None
                        
                except Exception as img_err:
                    logger.warning(f"Image extraction failed for page {page_num}: {img_err}")
            
            doc.close()
            
            # Save Word document
            word_doc.save(output_docx_path)
            logger.info(f"Created Word document: {output_docx_path}")
            
            # Save text file if requested
            txt_path = None
            if output_txt_path:
                with open(output_txt_path, 'w', encoding='utf-8') as f:
                    f.write('\n\n'.join(full_text))
                txt_path = output_txt_path
                logger.info(f"Created text file: {output_txt_path}")
            
            return output_docx_path, txt_path, None
            
        except Exception as e:
            logger.error(f"PDF to Word conversion failed: {e}")
            raise Exception(f"PDF Word ಪರಿವರ್ತನೆ ವಿಫಲ: {str(e)}")
