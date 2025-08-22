"""
Unified PDF to Word converter for Kannada documents.
Combines the best features from all modules into one clean interface.
"""

import os
import logging
import unicodedata
from typing import Tuple, Optional, List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber
try:
    from pdf2image import convert_from_path
    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False
    convert_from_path = None
from PIL import Image
import fitz  # PyMuPDF

from .legacy_kannada import (
    post_process_kannada_text,
    detect_legacy_encoding,
    is_kannada_text,
    normalize_unicode
)
from .kannada_image_preprocessor import preprocess_kannada_image

# Conditional imports for optional features
try:
    from .ocr_docai import docai_extract
    HAS_DOCAI = True
except ImportError:
    HAS_DOCAI = False
    docai_extract = None

try:
    from .ocr_vision import rasterize_pdf_to_images, vision_fulltext_on_image
    HAS_VISION = True
except ImportError:
    HAS_VISION = False
    rasterize_pdf_to_images = None
    vision_fulltext_on_image = None

try:
    from .gcs_io import upload_to_gcs, signed_url
    HAS_GCS = True
except ImportError:
    HAS_GCS = False
    upload_to_gcs = None
    signed_url = None

from .docx_builder import DocxBuilder

# Fallback converter for when poppler is not available
try:
    from .fallback_converter import FallbackPDFConverter
    HAS_FALLBACK = True
except ImportError:
    HAS_FALLBACK = False
    FallbackPDFConverter = None

logger = logging.getLogger(__name__)

class WatermarkFilter:
    """Detect and filter repeated headers/footers/watermarks across pages."""
    def __init__(self, min_pages_ratio: float = 0.6, y_margin_ratio: float = 0.12):
        self.min_pages_ratio = min_pages_ratio
        self.y_margin_ratio = y_margin_ratio  # Top/bottom band where headers/footers usually appear
        self.counter: Dict[str, int] = {}
        self.global_counter: Dict[str, int] = {}  # Track all repeated text, not just margins
        self.pages_seen = 0
        # Common watermark keywords
        self.watermark_keywords = ("scanned by", "confidential", "draft", "sample", "copy", "duplicate", "watermark")

    def next_page(self):
        self.pages_seen += 1

    @staticmethod
    def _normalize_text(text: str) -> str:
        t = (text or '').strip()
        t = t.replace('\u200d', '').replace('\u200c', '')  # remove ZWJ/ZWNJ
        t = ' '.join(t.split())
        return t.lower()

    def observe(self, page_h: float, bbox: List[float] | tuple, text: str):
        """Record a line candidate if it's in margin bands."""
        if not text or not bbox or page_h <= 0:
            return
        key = self._normalize_text(text)
        if not key or len(key) < 3:
            return

        # Track all text globally for repeated content detection
        self.global_counter[key] = self.global_counter.get(key, 0) + 1

        # Track margin text specifically
        y_top = bbox[1]
        y_bottom = bbox[3]
        band = self.y_margin_ratio * page_h
        if y_top <= band or (page_h - y_bottom) <= band:
            self.counter[key] = self.counter.get(key, 0) + 1

    def is_watermark(self, text: str) -> bool:
        if not text or not self.pages_seen:
            return False
        key = self._normalize_text(text)
        if not key:
            return False

        # Check for common watermark keywords
        if any(keyword in key for keyword in self.watermark_keywords):
            return True

        # Check margin-based repetition
        margin_freq = self.counter.get(key, 0)
        if self.pages_seen and (margin_freq / self.pages_seen) >= self.min_pages_ratio:
            return True

        # Check global repetition with higher threshold
        global_freq = self.global_counter.get(key, 0)
        if self.pages_seen and (global_freq / self.pages_seen) >= 0.7:
            return True

        return False

class UnifiedPDFConverter:
    """Single converter that handles all PDF types intelligently."""

    def __init__(self, use_google_vision: bool = False, debug_mode: bool = False,
                 normalize_font_sizes: bool = True, strip_watermark: bool = True,
                 fix_kannada_spacing_opt: bool = True):
        self.use_google_vision = use_google_vision
        self.debug_mode = debug_mode
        self.normalize_font_sizes = normalize_font_sizes
        self.strip_watermark = strip_watermark
        self.fix_kannada_spacing_opt = fix_kannada_spacing_opt

        self.warnings: List[str] = []

    def convert_pdf_to_word(
        self,
        input_pdf_path: str,
        output_docx_path: str,
        output_txt_path: Optional[str] = None,
        title: Optional[str] = None,
        author: Optional[str] = None,
        force_ocr: bool = False,
        mode: str = "auto",
        store_gcs: bool = False
    ) -> Tuple[str, Optional[str], Optional[Dict[str, str]]]:
        """
        Main conversion method with mode control:
          - auto: detect and choose best path
          - digital: force digital text extraction
          - scanned: force OCR with layout (DocAI -> Vision fallback)
          - fast: Google Drive OCR conversion (best-effort)
        """
        try:
            gcs_urls = None

            if mode == "digital":
                result = self._convert_digital_pdf(
                    input_pdf_path, output_docx_path, output_txt_path, title, author
                )
            elif mode == "scanned":
                result = self._convert_scanned_pdf_with_layout(
                    input_pdf_path, output_docx_path, output_txt_path, title, author
                )
            elif mode == "fast":
                result = self._convert_fast_drive(
                    input_pdf_path, output_docx_path, output_txt_path, title, author
                )
            else:  # auto
                forced = os.getenv("FORCE_PDF_MODE")
                if forced in ("digital", "scanned"):
                    pdf_type = forced
                    logger.info(f"FORCE_PDF_MODE active: {pdf_type}")
                else:
                    pdf_type = self._detect_pdf_type(input_pdf_path, force_ocr)
                logger.info(f"Detected PDF type: {pdf_type}")
                if pdf_type == "digital":
                    result = self._convert_digital_pdf(
                        input_pdf_path, output_docx_path, output_txt_path, title, author
                    )
                else:
                    result = self._convert_scanned_pdf_with_layout(
                        input_pdf_path, output_docx_path, output_txt_path, title, author
                    )

            # Upload to GCS if requested
            if store_gcs:
                gcs_urls = self._upload_to_gcs(result[0], result[1])

            return result[0], result[1], gcs_urls

        except Exception as e:
            error_msg = str(e).lower()
            # Check if it's a poppler-related error and use fallback
            if ("poppler" in error_msg or "unable to get page count" in error_msg or 
                "pdf2image" in error_msg) and HAS_FALLBACK:
                logger.warning(f"Poppler not available, using fallback converter: {e}")
                fallback = FallbackPDFConverter(debug_mode=False)
                result = fallback.convert_pdf_to_word_simple(
                    input_pdf_path, output_docx_path, output_txt_path, title, author
                )
                
                # Upload to GCS if requested
                gcs_urls = None
                if store_gcs:
                    gcs_urls = self._upload_to_gcs(result[0], result[1])
                
                return result[0], result[1], gcs_urls
            else:
                # Re-raise the original error
                raise

        except Exception as e:
            logger.error(f"Conversion failed: {e}")
            raise

    def _detect_pdf_type(self, pdf_path: str, force_ocr: bool = False) -> str:
        """Detect if PDF is digital or scanned."""
        if force_ocr:
            return "scanned"

        try:
            # Try to extract text
            with pdfplumber.open(pdf_path) as pdf:
                total_text = ""
                # Sample up to 10 pages across the document (no arbitrary low limits)
                page_count = len(pdf.pages)
                if page_count == 0:
                    return "scanned"
                sample_indices = list(range(min(page_count, 10)))  # first N pages
                # If more than 10 pages, also include last page for coverage
                if page_count > 10:
                    sample_indices[-1] = page_count - 1
                for idx in sample_indices:
                    try:
                        page = pdf.pages[idx]
                        text = page.extract_text() or ""
                        total_text += text
                    except Exception:
                        continue

                # Check if we have meaningful Kannada text
                if len(total_text.strip()) > 50 and is_kannada_text(total_text):
                    return "digital"
                else:
                    return "scanned"

        except Exception:
            return "scanned"

    def _convert_digital_pdf(
        self, input_pdf_path: str, output_docx_path: str,
        output_txt_path: Optional[str], title: Optional[str], author: Optional[str]
    ) -> Tuple[str, Optional[str]]:
        """Convert digital PDF preserving layout, fonts, and images with column and watermark handling."""
        builder = DocxBuilder(title=title, author=author)
        full_text = []
        suppressed_lines = 0
        sizes_observed: List[float] = []

        try:
            with fitz.open(input_pdf_path) as pdf:
                wm = WatermarkFilter()

                # First pass: detect watermarks/headers/footers and columns
                multi_column_detected = False
                for page in pdf:
                    wm.next_page()
                    page_dict = page.get_text("dict")
                    page_h = float(page.rect.height)
                    for block in page_dict.get("blocks", []):
                        if block.get("type", 0) != 0:
                            continue
                        for line in block.get("lines", []):
                            line_text = ''.join([span.get('text', '') for span in line.get('spans', [])]).strip()
                            bbox = line.get('bbox', [0, 0, 0, 0])
                            wm.observe(page_h, bbox, line_text)
                    # Simple column detection: check distribution of text block x centers
                    x_centers = []
                    for block in page_dict.get("blocks", []):
                        if block.get("type", 0) == 0:
                            bbox = block.get('bbox', [0, 0, 0, 0])
                            x_centers.append((bbox[0] + bbox[2]) / 2.0)
                    if len(x_centers) > 10:
                        x_centers_sorted = sorted(x_centers)
                        # Compute biggest gap between adjacent centers
                        gaps = [(x_centers_sorted[i+1] - x_centers_sorted[i]) for i in range(len(x_centers_sorted)-1)]
                        if gaps and max(gaps) > 72 * 1.0:  # >1 inch gap suggests two columns
                            multi_column_detected = True
                # Apply columns if detected
                if multi_column_detected:
                    builder.set_columns(2)

                # Second pass: build document content in reading order
                for page_index, page in enumerate(pdf):
                    page_dict = page.get_text("dict")
                    page_h = float(page.rect.height)
                    blocks = page_dict.get("blocks", [])

                    # If two columns detected, split blocks by left x
                    if multi_column_detected:
                        left_col, right_col = [], []
                        # Determine page mid x
                        mid_x = page.rect.width / 2.0
                        for b in blocks:
                            if b.get("type", 0) != 0:
                                continue
                            x0, _, x1, _ = b.get('bbox', [0, 0, 0, 0])
                            cx = (x0 + x1) / 2.0
                            (left_col if cx <= mid_x else right_col).append(b)
                        cols = [sorted(left_col, key=lambda bb: bb.get('bbox', [0, 0, 0, 0])[1]),
                                sorted(right_col, key=lambda bb: bb.get('bbox', [0, 0, 0, 0])[1])]
                    else:
                        cols = [sorted([b for b in blocks if b.get("type", 0) == 0],
                                        key=lambda bb: bb.get('bbox', [0, 0, 0, 0])[1])]

                    # Emit text lines
                    for col_blocks in cols:
                        for b in col_blocks:
                            for line in b.get("lines", []):
                                # Aggregate spans into a single paragraph line
                                spans = []
                                line_text_parts = []
                                bbox = line.get('bbox', [0, 0, 0, 0])
                                # Skip repeated watermark/header/footer lines
                                preview_text = ''.join([s.get('text', '') for s in line.get('spans', [])]).strip()
                                if self.strip_watermark and wm.is_watermark(preview_text):
                                    continue
                                for span in line.get("spans", []):
                                    raw = span.get("text", "")
                                    if not raw:
                                        continue
                                    # Kannada-aware cleanup
                                    is_legacy = detect_legacy_encoding(raw)
                                    processed = post_process_kannada_text(raw, is_legacy)
                                    # Optional Kannada spacing fix
                                    if getattr(self, 'fix_kannada_spacing_opt', True):
                                        try:
                                            from .kannada_text_post import fix_kannada_spacing
                                            processed = fix_kannada_spacing(processed)
                                        except Exception:
                                            pass
                                    processed = self._normalize_kannada_text(processed)

                                    # Build span with styling
                                    fname = (span.get("font") or "").lower()
                                    size = span.get("size", 12)

                                    def _norm_size(pt):
                                        for s in (10, 11, 12, 14, 16, 18, 22):
                                            if pt <= s:
                                                return s
                                        return 24

                                    span_data = {
                                        'text': processed,
                                        'bold': "bold" in fname,
                                        'italic': "italic" in fname or "oblique" in fname,
                                        'size_pt': _norm_size(float(size)),
                                        'font': 'Noto Sans Kannada'
                                    }
                                    spans.append(span_data)
                                    line_text_parts.append(processed)
                                if spans:
                                    # Add small space before paragraphs not at top
                                    para_meta = {'_para': {'space_before_pt': 2}}
                                    builder.add_paragraph([para_meta] + spans)
                                    full_text.append(''.join(line_text_parts))

                    # Add images where available (preserve approximate width)
                    for b in blocks:
                        if b.get("type", 0) == 1:  # image block
                            xref = b.get("image")
                            bbox = b.get("bbox", [0, 0, 0, 0])
                            width_in = ((bbox[2] - bbox[0]) / 72.0) * 0.9
                            if xref:
                                try:
                                    pix = fitz.Pixmap(page.parent, xref)
                                    if pix.alpha:
                                        pix = fitz.Pixmap(fitz.csRGB, pix)
                                    img_bytes = pix.tobytes("png")
                                    builder.add_image(img_bytes, max(0.8, width_in))
                                except Exception as e:
                                    logger.debug(f"Failed to embed image xref {xref}: {e}")

                    # Page break between pages
                    if page_index < len(pdf) - 1:
                        builder.page_break()

            # Save files
            builder.save(output_docx_path)

            if output_txt_path:
                with open(output_txt_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(full_text))

            return output_docx_path, output_txt_path

        except Exception as e:
            logger.error(f"Digital PDF conversion failed: {e}")
            raise

    def _convert_scanned_pdf_with_layout(
        self, input_pdf_path: str, output_docx_path: str,
        output_txt_path: Optional[str], title: Optional[str], author: Optional[str]
    ) -> Tuple[str, Optional[str]]:
        """Convert scanned PDF using Document AI with Vision fallback, preserving layout and removing watermarks."""
        builder = DocxBuilder(title=title, author=author)
        all_text = []

        try:
            # Get language hints
            lang_hints = [s.strip() for s in os.getenv('LANGUAGE_HINTS', 'kn,en').split(',')]

            # Try Document AI first
            try:
                if not HAS_DOCAI:
                    raise ImportError("Document AI not available")

                # Chunking for Document AI to avoid size/page limits
                import fitz
                all_pages = []
                with fitz.open(input_pdf_path) as full_pdf:
                    total_pages = len(full_pdf)

                    max_pages_per_chunk = int(os.getenv('DOCAI_MAX_PAGES_PER_CHUNK', '15'))
                    max_bytes_per_chunk = int(os.getenv('DOCAI_MAX_BYTES_PER_CHUNK', str(8 * 1024 * 1024)))
                    cur = 0
                    while cur < total_pages:
                        chunk_doc = fitz.open()
                        start = cur
                        while cur < total_pages and (cur - start) < max_pages_per_chunk:
                            chunk_doc.insert_pdf(full_pdf, from_page=cur, to_page=cur)
                            cur += 1
                            # Check size
                            try:
                                if len(chunk_doc.tobytes()) >= max_bytes_per_chunk:
                                    # remove last page and stop this chunk
                                    cur -= 1
                                    chunk_doc = fitz.open()
                                    if cur > start:
                                        chunk_doc.insert_pdf(full_pdf, from_page=start, to_page=cur-1)
                                    break
                            except Exception:
                                break
                        chunk_bytes = chunk_doc.tobytes()
                        chunk_pages = docai_extract(chunk_bytes, lang_hints)
                        all_pages.extend(chunk_pages)
                pages = all_pages
                # Quality gate: fallback to Vision if DocAI produced too little text
                try:
                    total_chars = 0
                    for p in pages:
                        for b in getattr(p, 'blocks', []):
                            total_chars += len(getattr(b, 'text', '') or '')
                    min_chars = int(os.getenv('DOCAI_MIN_CHARS', '200'))
                    if total_chars < min_chars:
                        raise RuntimeError(f"DocAI low content: {total_chars} chars (<{min_chars})")
                except Exception as _q:
                    raise
                logger.info(f"Document AI extracted {len(pages)} pages across chunks")

                # Build watermark model using page margins with PyMuPDF bboxes
                wm = WatermarkFilter()
                try:
                    with fitz.open(input_pdf_path) as pdf:
                        for page in pdf:
                            wm.next_page()
                            page_h = float(page.rect.height)
                            page_dict = page.get_text("dict")
                            for block in page_dict.get("blocks", []):
                                if block.get("type", 0) != 0:
                                    continue
                                for line in block.get("lines", []):
                                    line_text = ''.join([span.get('text', '') for span in line.get('spans', [])]).strip()
                                    bbox = line.get('bbox', [0, 0, 0, 0])
                                    wm.observe(page_h, bbox, line_text)
                except Exception:
                    pass

                for page_idx, page in enumerate(pages):
                    # Add text blocks with simple reading order
                    for block in page.blocks:
                        for paragraph in block.paragraphs:
                            for line in paragraph.lines:
                                # Skip watermark-like lines using text match
                                if wm.is_watermark(line.text):
                                    continue
                                spans = []
                                line_text = ""

                                for span in line.spans:
                                    processed = self._normalize_kannada_text(span.text)
                                    # Use stable body font size to avoid jitter
                                    span_data = {
                                        'text': processed,
                                        'bold': False,
                                        'italic': False,
                                        'size_pt': 12.0,
                                        'font': 'Noto Sans Kannada'
                                    }
                                    spans.append(span_data)
                                    line_text += processed

                                if spans:
                                    builder.add_paragraph(spans)
                                    all_text.append(line_text)

                    # Add tables
                    for table in page.tables:
                        if table:
                            builder.add_table(table)

                    # Page break
                    if page_idx < len(pages) - 1:
                        builder.page_break()

                # Extract and insert images separately
                self._extract_pdf_images(input_pdf_path, builder)

            except Exception as e:
                logger.warning(f"Document AI failed, falling back to Vision API: {e}")

                if not HAS_VISION:
                    raise ImportError("Vision API not available")

                # Vision API fallback
                images = rasterize_pdf_to_images(input_pdf_path)

                for idx, img in enumerate(images):
                    try:
                        # Preprocess for better OCR
                        processed_img = preprocess_kannada_image(img)
                        pil_img = Image.fromarray(processed_img)

                        # Extract text spans
                        line_spans = vision_fulltext_on_image(pil_img, lang_hints)

                        page_text_parts = []
                        for line_span in line_spans:
                            # Filter obvious watermark-like ASCII-heavy lines if enabled
                            if self.strip_watermark:
                                txt = (line_span.text or "").strip()
                                ascii_ratio = sum(1 for ch in txt if (ch.isascii() and ch.isalpha()) or ch == ' ') / max(1, len(txt))
                                if ascii_ratio > 0.9 and len(txt) > 6:
                                    continue

                            processed = self._normalize_kannada_text(line_span.text)
                            if getattr(self, 'fix_kannada_spacing_opt', True):
                                try:
                                    from .kannada_text_post import fix_kannada_spacing
                                    processed = fix_kannada_spacing(processed)
                                except Exception:
                                    pass

                            span_data = {
                                'text': processed,
                                'bold': False,
                                'italic': False,
                                'size_pt': 12.0,
                                'font': 'Noto Sans Kannada'
                            }
                            builder.add_paragraph([span_data])
                            page_text_parts.append(processed)

                        if page_text_parts:
                            all_text.extend(page_text_parts)

                        # Add original page image for layout context (optional)
                        import io
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        builder.add_image(img_byte_arr.getvalue(), 5.0)

                        if idx < len(images) - 1:
                            builder.page_break()

                    except Exception as page_e:
                        logger.warning(f"Vision OCR failed on page {idx + 1}: {page_e}")

            # Save outputs
            builder.save(output_docx_path)

            if output_txt_path:
                with open(output_txt_path, 'w', encoding='utf-8') as f:
                    f.write('\n\n'.join(all_text))

            return output_docx_path, output_txt_path

        except Exception as e:
            logger.error(f"Scanned PDF conversion failed: {e}")
            raise

    def _convert_fast_drive(
        self, input_pdf_path: str, output_docx_path: str,
        output_txt_path: Optional[str], title: Optional[str], author: Optional[str]
    ) -> Tuple[str, Optional[str]]:
        """Fast mode: Use Google Drive OCR conversion."""
        try:
            try:
                from googleapiclient.discovery import build
                from googleapiclient.http import MediaFileUpload
                from google.oauth2 import service_account
            except ImportError as e:
                raise ImportError(f"Google API packages not installed: {e}. Run: pip install google-api-python-client google-auth")

            scopes = [
                'https://www.googleapis.com/auth/drive',
                'https://www.googleapis.com/auth/drive.file',
            ]
            cred_path = os.getenv('GOOGLE_APPLICATION_CREDENTIALS')
            if not cred_path or not os.path.exists(cred_path):
                raise RuntimeError('Service account credentials not found for Drive fast mode')

            creds = service_account.Credentials.from_service_account_file(cred_path, scopes=scopes)
            drive = build('drive', 'v3', credentials=creds, cache_discovery=False)

            # Upload PDF to Drive as Google Doc
            media = MediaFileUpload(input_pdf_path, mimetype='application/pdf', resumable=False)
            file_metadata = {
                'name': os.path.basename(input_pdf_path),
                'mimeType': 'application/vnd.google-apps.document'
            }
            created = drive.files().create(body=file_metadata, media_body=media, fields='id').execute()
            file_id = created['id']

            # Export as DOCX and TXT
            docx_bytes = drive.files().export(
                fileId=file_id,
                mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            ).execute()

            with open(output_docx_path, 'wb') as f:
                f.write(docx_bytes)

            if output_txt_path:
                txt_bytes = drive.files().export(fileId=file_id, mimeType='text/plain').execute()
                with open(output_txt_path, 'wb') as f:
                    f.write(txt_bytes)

            # Cleanup
            try:
                drive.files().delete(fileId=file_id).execute()
            except Exception:
                pass

            return output_docx_path, output_txt_path

        except Exception as e:
            logger.error(f"Fast Drive OCR conversion failed: {e}")
            raise

    def _normalize_kannada_text(self, text: str) -> str:
        """NFC-only normalization plus Kannada-safe spacing fixes.
        Avoid stripping ZWJ/ZWNJ or virama; only remove ASCII spaces inside grapheme clusters.
        """
        if not text:
            return text

        import re
        # NFC normalization only
        text = unicodedata.normalize('NFC', text)

        # Unicode-safe spacing fix using dedicated module
        try:
            from .kannada_text_post import fix_kannada_spacing
            text = fix_kannada_spacing(text)
        except Exception:
            # Fallback minimal rules: remove spaces around halant and matras
            text = re.sub(r'([\u0C80-\u0CFF])\s+([\u0CBE-\u0CD6\u0CCD\u200C\u200D])', r'\1\2', text)
            text = re.sub(r'([\u0CBE-\u0CD6\u0CCD\u200C\u200D])\s+([\u0C80-\u0CFF])', r'\1\2', text)
            text = re.sub(r'\s+', ' ', text).strip()

        return text

    def _extract_pdf_images(self, pdf_path: str, builder: DocxBuilder):
        """Extract images from PDF and add to document."""
        try:
            with fitz.open(pdf_path) as pdf:
                for page in pdf:
                    imgs = page.get_images(full=True)
                    for (xref, *_) in imgs:
                        try:
                            pix = fitz.Pixmap(pdf, xref)
                            if pix.alpha:
                                pix = fitz.Pixmap(fitz.csRGB, pix)
                            img_bytes = pix.tobytes("png")
                            builder.add_image(img_bytes, 4.0)
                        except Exception:
                            continue
        except Exception as e:
            logger.debug(f"Image extraction failed: {e}")

    def _upload_to_gcs(self, docx_path: str, txt_path: Optional[str]) -> Dict[str, str]:
        """Upload files to GCS and return signed URLs."""
        if not HAS_GCS:
            logger.warning("GCS not available, skipping upload")
            return {}

        try:
            import uuid
            from datetime import datetime

            bucket_name = os.getenv("GCS_OUTPUT_BUCKET")
            if not bucket_name:
                raise ValueError("GCS_OUTPUT_BUCKET not configured")

            unique_id = str(uuid.uuid4())[:8]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            urls = {}

            # Upload DOCX
            docx_key = f"docx/{timestamp}_{unique_id}.docx"
            docx_gs_uri = upload_to_gcs(docx_path, bucket_name, docx_key)
            urls['docx'] = signed_url(docx_gs_uri, minutes=60)

            # Upload TXT if available
            if txt_path and os.path.exists(txt_path):
                txt_key = f"txt/{timestamp}_{unique_id}.txt"
                txt_gs_uri = upload_to_gcs(txt_path, bucket_name, txt_key)
                urls['txt'] = signed_url(txt_gs_uri, minutes=60)

            return urls

        except Exception as e:
            logger.error(f"GCS upload failed: {e}")
            return {}

# Convenience functions for backward compatibility
def convert_pdf_to_word(input_pdf_path: str, output_docx_path: str, **kwargs):
    """Legacy function for backward compatibility."""
    converter = UnifiedPDFConverter()
    result = converter.convert_pdf_to_word(input_pdf_path, output_docx_path, **kwargs)
    return result[0], result[1]

def ocr_pdf_to_word(input_pdf_path: str, output_docx_path: str, **kwargs):
    """Legacy function for backward compatibility."""
    converter = UnifiedPDFConverter(use_google_vision=kwargs.get('use_google', False))
    result = converter.convert_pdf_to_word(input_pdf_path, output_docx_path, force_ocr=True, **kwargs)
    return result[0], result[1]