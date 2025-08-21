import io
from typing import List, Dict, Any, Optional, Union
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import logging

logger = logging.getLogger(__name__)

class DocxBuilder:
    def __init__(self, page_width_in: float = 6.0, title: Optional[str] = None, author: Optional[str] = None):
        self.doc = Document()
        self.page_width_in = page_width_in
        
        # Set document properties
        if title:
            self.doc.core_properties.title = title
        if author:
            self.doc.core_properties.author = author
        
        # Set default font for better Kannada support
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'  # Fallback font
        try:
            # Try to set Kannada-friendly font
            font.name = 'Noto Sans Kannada'
        except:
            try:
                font.name = 'Tunga'  # Alternative Kannada font
            except:
                pass  # Use fallback if no Kannada fonts available
        font.size = Pt(12)
    
    def set_columns(self, num_columns: int = 1):
        """Set number of columns for the current section (applies document-wide unless new sections are added)."""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            section = self.doc.sections[-1]
            sectPr = section._sectPr
            cols = sectPr.xpath('./w:cols')
            if cols:
                cols[0].set(qn('w:num'), str(num_columns))
            else:
                cols_elm = OxmlElement('w:cols')
                cols_elm.set(qn('w:num'), str(num_columns))
                sectPr.append(cols_elm)
        except Exception as e:
            logger.debug(f"Failed to set columns: {e}")
    
    def add_paragraph(self, spans: List[Dict[str, Any]]):
        """Add paragraph with styled spans"""
        p = self.doc.add_paragraph()
        
        # Optional paragraph-level props via special key on first span
        if spans and isinstance(spans[0], dict) and spans[0].get('_para'):
            props = spans[0]['_para']
            # Alignment
            align = props.get('align')
            if align:
                if isinstance(align, int):
                    p.alignment = align
                elif isinstance(align, str):
                    align_l = align.lower()
                    if align_l == 'center':
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif align_l == 'right':
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    else:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # Indentation and spacing (points)
            if 'left_indent_pt' in props:
                p.paragraph_format.left_indent = Pt(props['left_indent_pt'])
            if 'space_before_pt' in props:
                p.paragraph_format.space_before = Pt(props['space_before_pt'])
            if 'space_after_pt' in props:
                p.paragraph_format.space_after = Pt(props['space_after_pt'])
            # Optional paragraph style
            style_name = props.get('style')
            if style_name:
                try:
                    p.style = style_name
                except Exception as e:
                    logger.debug(f"Unknown style '{style_name}': {e}")
        
        # Line-level size smoothing: snap small differences to dominant size
        sizes = [float(s.get('size_pt', 12)) for s in spans if isinstance(s, dict) and not s.get('_para')]
        dominant = None
        if sizes:
            # Round to nearest 0.5pt for stability
            rounded = [round(v * 2) / 2.0 for v in sizes]
            # Pick mode-ish value
            freq = {}
            for v in rounded:
                freq[v] = freq.get(v, 0) + 1
            dominant = max(freq.items(), key=lambda kv: kv[1])[0]

        for span in spans:
            text = span.get('text', '')
            # Skip internal control dicts
            if span.get('_para'):
                continue
            bold = span.get('bold', False)
            italic = span.get('italic', False)
            size_pt = float(span.get('size_pt', 12))
            font_name = span.get('font', 'Noto Sans Kannada')
            # Snap to dominant if within 0.75pt
            if dominant is not None and abs((round(size_pt*2)/2.0) - dominant) < 0.75:
                size_pt = dominant

            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size_pt)
            try:
                run.font.name = font_name
            except:
                run.font.name = 'Calibri'  # Fallback
    
    def add_table(self, rows: Union[List[List[str]], List[List[Dict[str, Any]]]]):
        """Add table with row data - supports both string and formatted cell content"""
        if not rows or not rows[0]:
            return
        
        # Determine max columns
        max_cols = max(len(row) for row in rows)
        
        table = self.doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_content in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    
                    # Handle both string and formatted content
                    if isinstance(cell_content, str):
                        cell.text = cell_content
                    elif isinstance(cell_content, dict):
                        # Formatted cell content
                        text = cell_content.get('text', '')
                        bold = cell_content.get('bold', False)
                        italic = cell_content.get('italic', False)
                        
                        cell_paragraph = cell.paragraphs[0]
                        run = cell_paragraph.runs[0] if cell_paragraph.runs else cell_paragraph.add_run()
                        run.text = text
                        run.bold = bold
                        run.italic = italic
                    else:
                        cell.text = str(cell_content)
    
    def add_image(self, img_bytes: bytes, width_in: Optional[float] = None, alignment: str = 'left'):
        """Add image with optional width constraint and alignment"""
        try:
            if width_in is None:
                width_in = min(self.page_width_in, 5.0)  # Default constraint
            
            # Ensure width is reasonable
            width_in = max(0.5, min(width_in, self.page_width_in))
            
            img_stream = io.BytesIO(img_bytes)
            
            # Create paragraph for image alignment
            p = self.doc.add_paragraph()
            run = p.add_run()
            run.add_picture(img_stream, width=Inches(width_in))
            
            # Set alignment
            if alignment.lower() == 'center':
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment.lower() == 'right':
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
        except Exception as e:
            logger.warning(f"Failed to add image: {e}")
            # Add placeholder text instead
            p = self.doc.add_paragraph("[Image could not be inserted]")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def add_heading(self, text: str, level: int = 1):
        """Add a heading with specified level"""
        heading = self.doc.add_heading(text, level=level)
        # Set Kannada font for heading
        for run in heading.runs:
            try:
                run.font.name = 'Noto Sans Kannada'
            except:
                try:
                    run.font.name = 'Tunga'
                except:
                    pass
    
    def page_break(self):
        """Add page break"""
        self.doc.add_page_break()
    
    def save(self, path: str):
        """Save document to file"""
        self.doc.save(path)
        logger.info(f"Document saved to {path}")

